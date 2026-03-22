import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date

st.set_page_config(page_title="History Taking — Dr. Hany Elhennawy", page_icon="🧠", layout="wide")

st.markdown("""
<style>
    .main-title { font-size: 26px; font-weight: 700; margin-bottom: 2px; color: #1A5CB8; }
    .sub-title { color: #888; font-size: 13px; margin-bottom: 24px; }
    .sec-header { font-size: 17px; font-weight: 700; color: #1A5CB8; margin-top: 18px; margin-bottom: 6px;
                  border-bottom: 2px solid #1A5CB8; padding-bottom: 4px; }
    .ar { direction: rtl; text-align: right; font-size: 14px; color: #333; }
    .bilingual { display: flex; justify-content: space-between; font-size: 13px; color: #555; margin-bottom: 2px; }
    .step-box { background: #f0f6ff; border-radius: 8px; padding: 10px 16px;
                margin-bottom: 18px; font-size: 13px; color: #1A5CB8; font-weight: 500; }
</style>
""", unsafe_allow_html=True)

# ── CONSTANTS ──
RECIPIENT_EMAIL = "yusuf.a.abdelatti@gmail.com"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "logo.png")
CLINIC_BLUE = RGBColor(0x1A, 0x5C, 0xB8)
NAVY = RGBColor(0x1B, 0x2A, 0x4A)

DOCTOR = {
    "name": "Dr. Hany Elhennawy",
    "title1": "Consultant of Neuro-Psychiatry",
    "title2": "Aviation Medical Council — Faculty of Medicine, 6th October University",
    "title3": "MD of Neuroscience Research, Karolinska Institute — Sweden",
    "title4": "Member of I.S.N.R",
    "address": "16 Hesham Labib St., off Makram Ebeid St. Ext., next to Mobilia Saad Mohamed Saad",
    "phone": "+20 1000756200",
}

# ── SIDEBAR ──
with st.sidebar:
    st.header("⚙️ Settings")
    groq_key = st.text_input("Groq API Key", type="password", placeholder="gsk_...")
    st.caption("Get free key at [console.groq.com](https://console.groq.com)")
    st.divider()
    gmail_user = st.text_input("Gmail Address (sender)", placeholder="yourname@gmail.com")
    gmail_pass = st.text_input("Gmail App Password", type="password", placeholder="xxxx xxxx xxxx xxxx")
    st.caption("Use a Gmail App Password — [how to get one](https://support.google.com/accounts/answer/185833)")
    st.divider()
    st.caption("**History by (Psychologist name):**")
    history_by = st.text_input("Psychologist Name", value=st.session_state.get("history_by", ""))

st.markdown('<div class="main-title">🧠 History Taking Sheet</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Dr. Hany Elhennawy Clinic — Neuro-Psychiatry & Neurofeedback</div>', unsafe_allow_html=True)

# ── STEP 0: CHOOSE TYPE ──
if "sheet_type" not in st.session_state:
    st.markdown("### Choose History Sheet Type / اختر نوع الاستمارة")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("👤 Adult / بالغ", use_container_width=True, type="primary"):
            st.session_state.sheet_type = "adult"
            st.session_state.step = 1
            st.rerun()
    with col2:
        if st.button("👶 Child / طفل", use_container_width=True):
            st.session_state.sheet_type = "child"
            st.session_state.step = 1
            st.rerun()
    st.stop()

sheet_type = st.session_state.sheet_type
step = st.session_state.get("step", 1)

# Step indicator
if sheet_type == "adult":
    steps = ["Personal Details", "Family Details", "Marriage Details", "Siblings",
             "Complaints & HPI", "Drug / Past / Family History", "Investigations & Surgeries", "Extra Clinical"]
else:
    steps = ["Personal Details", "Family Details", "Siblings",
             "Complaints & HPI", "Past / Family History", "Investigations & Surgeries", "Child Checklist"]

total = len(steps)
st.markdown(f'<div class="step-box">Step {step} of {total}: {steps[step-1]}</div>', unsafe_allow_html=True)
st.progress(step / total)

def bilingual(en, ar):
    st.markdown(f'<div class="bilingual"><span>{en}</span><span class="ar">{ar}</span></div>', unsafe_allow_html=True)

def sec(en, ar=""):
    st.markdown(f'<div class="sec-header">{en}{" / " + ar if ar else ""}</div>', unsafe_allow_html=True)

def nav(prev_step, next_step=None, final=False):
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← Back / رجوع") and prev_step:
            st.session_state.step = prev_step; st.rerun()
    with col2:
        if final:
            if st.button("✦ Generate Report / إنشاء التقرير", type="primary"):
                if not groq_key:
                    st.error("Please enter your Groq API key in the sidebar.")
                elif not gmail_user or not gmail_pass:
                    st.error("Please enter your Gmail credentials in the sidebar.")
                else:
                    st.session_state.generate = True; st.rerun()
        else:
            if st.button("Next / التالي →", type="primary"):
                st.session_state.step = next_step; st.rerun()

def ti(label_en, label_ar, key, placeholder=""):
    bilingual(label_en, label_ar)
    return st.text_input("", key=key, placeholder=placeholder, label_visibility="collapsed")

def ta(label_en, label_ar, key, height=100):
    bilingual(label_en, label_ar)
    return st.text_area("", key=key, height=height, label_visibility="collapsed")

def rb(label_en, label_ar, opts, key):
    bilingual(label_en, label_ar)
    return st.radio("", opts, key=key, horizontal=True, label_visibility="collapsed")

# ════════════════════════════════════════════════════════
#  ADULT SHEET
# ════════════════════════════════════════════════════════
if sheet_type == "adult":

    if step == 1:
        sec("Personal Details", "البيانات الشخصية")
        c1, c2 = st.columns(2)
        with c1:
            ti("Full Name", "الاسم", "name")
            ti("Age", "السن", "age")
            ti("Gender", "النوع", "gender")
            ti("Occupation / Study", "الوظيفة / الدراسة", "occupation")
            ti("Education Level", "المستوى التعليمي", "education")
        with c2:
            ti("Social Status", "الحالة الاجتماعية", "social_status")
            ti("Hobbies", "الهوايات", "hobbies")
            ti("Smoking", "التدخين", "smoking")
            ti("Phone Number", "رقم الهاتف", "phone")
            ti("Referral Source", "مصدر الإحالة", "referral")
        ti("Taken Date", "تاريخ الجلسة", "taken_date", placeholder=str(date.today()))
        ti("History Type", "نوع التاريخ", "history_type")
        nav(None, 2)

    elif step == 2:
        sec("Family Details", "بيانات الأسرة")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Father / الأب**")
            ti("Father's Name", "اسم الأب", "father_name")
            ti("Father's Age", "سن الأب", "father_age")
            ti("Father's Occupation", "وظيفة الأب", "father_occ")
        with c2:
            st.markdown("**Mother / الأم**")
            ti("Mother's Name", "اسم الأم", "mother_name")
            ti("Mother's Age", "سن الأم", "mother_age")
            ti("Mother's Occupation", "وظيفة الأم", "mother_occ")
        ti("Consanguinity (relation between parents)", "صلة القرابة بين الأب والأم", "consanguinity")
        ti("Chronic Illness in Family", "مرض مزمن في الأسرة", "chronic_illness")
        nav(1, 3)

    elif step == 3:
        sec("Marriage Details", "بيانات الزواج")
        ti("Spouse Name", "اسم الزوج / الزوجة", "spouse_name")
        ti("Spouse Age", "سن الزوج / الزوجة", "spouse_age")
        ti("Spouse Occupation", "وظيفة الزوج / الزوجة", "spouse_occ")
        ti("Duration of Marriage", "فترة الزواج", "marriage_duration")
        ti("Engagement Period", "فترة الخطوبة", "engagement")
        rb("Was there a contract before marriage? (كتب كتاب)", "كتب كتاب", ["Yes / نعم", "No / لا", "N/A"], "katb_ketab")
        ti("Relationship before marriage", "العلاقة قبل الزواج", "pre_marriage_rel")
        ti("Number of Children", "عدد الأبناء", "num_children")
        nav(2, 4)

    elif step == 4:
        sec("Brothers and Sisters", "الإخوة والأخوات")
        siblings = []
        for i in range(1, 5):
            st.markdown(f"**Sibling {i} / الأخ/الأخت {i}**")
            c1, c2, c3, c4 = st.columns(4)
            with c1: g = st.text_input("Gender/النوع", key=f"sib_gender_{i}", label_visibility="collapsed", placeholder="Gender/النوع")
            with c2: n = st.text_input("Name/الاسم", key=f"sib_name_{i}", label_visibility="collapsed", placeholder="Name/الاسم")
            with c3: a = st.text_input("Age/السن", key=f"sib_age_{i}", label_visibility="collapsed", placeholder="Age/السن")
            with c4: ed = st.text_input("Education/التعليم", key=f"sib_edu_{i}", label_visibility="collapsed", placeholder="Education/التعليم")
            notes = st.text_input("Notes/ملاحظات", key=f"sib_notes_{i}", label_visibility="collapsed", placeholder="Notes / ملاحظات")
            siblings.append({"gender": g, "name": n, "age": a, "edu": ed, "notes": notes})
        st.session_state["siblings"] = siblings
        nav(3, 5)

    elif step == 5:
        sec("Complaints & History of Presenting Illness", "الشكاوى وتاريخ المرض الحالي")
        ta("Chief Complaints (C/O)", "الشكاوى الرئيسية", "complaints", 120)
        ti("Onset / Since when?", "متى بدأت الأعراض؟", "onset")
        ta("History of Presenting Illness (HPI)", "تاريخ المرض الحالي", "hpi", 250)
        nav(4, 6)

    elif step == 6:
        sec("Drug History", "تاريخ الأدوية")
        ta("Current and past medications", "الأدوية الحالية والسابقة", "drug_history", 100)
        sec("Past History", "التاريخ المرضي السابق")
        ta("Previous illnesses / hospitalizations", "الأمراض السابقة / دخول المستشفى", "past_history", 100)
        sec("Family History", "التاريخ العائلي")
        ta("Psychiatric or neurological illness in family", "أمراض نفسية أو عصبية في الأسرة", "family_history", 100)
        nav(5, 7)

    elif step == 7:
        sec("Investigations", "الفحوصات")
        ta("Lab work, EEG, MRI, CT, etc.", "تحاليل، رسم مخ، رنين، أشعة مقطعية...", "investigations", 100)
        sec("Operations and Surgeries", "العمليات والجراحات")
        ta("Previous surgeries", "العمليات الجراحية السابقة", "surgeries", 100)
        nav(6, 8)

    elif step == 8:
        sec("Extra Clinical Assessment", "التقييم السريري الإضافي")
        rb("Sleep pattern / نمط النوم", "نمط النوم", ["Normal / طبيعي", "Insomnia / أرق", "Hypersomnia / نوم زيادة", "Disrupted / متقطع"], "sleep")
        rb("Appetite / الشهية", "الشهية", ["Normal / طبيعي", "Decreased / قلت", "Increased / زادت"], "appetite")
        rb("Suicidal ideation / أفكار انتحارية", "أفكار انتحارية", ["None / لا", "Passive / سلبية", "Active / نشطة"], "suicidal")
        rb("Substance use beyond tobacco / تعاطي مواد", "تعاطي مواد أخرى", ["None / لا", "Yes — specify below / نعم — حدد"], "substance")
        ta("Substance details if yes", "تفاصيل المواد إن وجدت", "substance_details", 60)
        ta("Additional notes / ملاحظات إضافية", "ملاحظات إضافية", "extra_notes", 100)
        nav(7, final=True)

# ════════════════════════════════════════════════════════
#  CHILD SHEET
# ════════════════════════════════════════════════════════
else:

    if step == 1:
        sec("Personal Details", "البيانات الشخصية")
        c1, c2 = st.columns(2)
        with c1:
            ti("Child's Full Name", "اسم الطفل", "name")
            ti("Age", "السن", "age")
            ti("Gender", "النوع", "gender")
            ti("School Name", "اسم المدرسة", "school_name")
            ti("Grade / Year", "الصف الدراسي", "grade")
        with c2:
            rb("Academic Performance / المستوى الدراسي", "المستوى الدراسي",
               ["Excellent / ممتاز", "Good / جيد", "Average / متوسط", "Weak / ضعيف"], "academic")
            ti("Who does child live with?", "الطفل يعيش مع", "lives_with")
            ti("Phone", "تليفون", "phone")
            ti("Taken Date", "تاريخ الجلسة", "taken_date", placeholder=str(date.today()))
        st.markdown("**Developmental Milestones / مراحل النمو**")
        c1, c2, c3 = st.columns(3)
        with c1:
            ti("Pregnancy / الحمل", "الحمل", "pregnancy")
            ti("Birth / الولادة", "الولادة", "birth")
            ti("Breastfeeding / الرضاعة", "الرضاعة", "breastfeeding")
            ti("Weaning / الفطام", "الفطام", "weaning")
        with c2:
            ti("Motor development / الحركة", "الحركة", "motor_dev")
            ti("Teething / التسنين", "التسنين", "teething")
            ti("Speech / الكلام", "الكلام", "speech")
            ti("Toilet training age", "سن تدريب دورة المياه", "toilet_training")
        with c3:
            ti("Vaccinations / التطعيمات", "التطعيمات", "vaccinations")
            ti("Immunization notes", "ملاحظات التطعيم", "vaccination_notes")
        ta("Developmental notes / ملاحظات النمو", "ملاحظات", "dev_notes", 80)
        nav(None, 2)

    elif step == 2:
        sec("Family Details", "بيانات الأسرة")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Father / الأب**")
            ti("Father's Name", "اسم الأب", "father_name")
            ti("Father's Age", "سن الأب", "father_age")
            ti("Father's Occupation", "وظيفة الأب", "father_occ")
            ti("Father hereditary illness", "مرض وراثي — الأب", "father_hereditary")
        with c2:
            st.markdown("**Mother / الأم**")
            ti("Mother's Name", "اسم الأم", "mother_name")
            ti("Mother's Age", "سن الأم", "mother_age")
            ti("Mother's Occupation", "وظيفة الأم", "mother_occ")
            ti("Mother hereditary illness", "مرض وراثي — الأم", "mother_hereditary")
        ti("Consanguinity between parents", "صلة القرابة بين الأب والأم", "consanguinity")
        ti("Parents relationship / علاقة الأب والأم ببعض", "علاقة الأب والأم", "parents_relation")
        nav(1, 3)

    elif step == 3:
        sec("Brothers and Sisters", "الإخوة والأخوات")
        siblings = []
        for i in range(1, 5):
            st.markdown(f"**Sibling {i} / الأخ/الأخت {i}**")
            c1, c2, c3, c4 = st.columns(4)
            with c1: g = st.text_input("Gender/النوع", key=f"sib_gender_{i}", label_visibility="collapsed", placeholder="Gender/النوع")
            with c2: n = st.text_input("Name/الاسم", key=f"sib_name_{i}", label_visibility="collapsed", placeholder="Name/الاسم")
            with c3: a = st.text_input("Age/السن", key=f"sib_age_{i}", label_visibility="collapsed", placeholder="Age/السن")
            with c4: ed = st.text_input("Education/التعليم", key=f"sib_edu_{i}", label_visibility="collapsed", placeholder="Education/التعليم")
            notes = st.text_input("Notes/ملاحظات", key=f"sib_notes_{i}", label_visibility="collapsed", placeholder="Notes / ملاحظات")
            siblings.append({"gender": g, "name": n, "age": a, "edu": ed, "notes": notes})
        st.session_state["siblings"] = siblings
        nav(2, 4)

    elif step == 4:
        sec("Complaints & History of Presenting Illness", "الشكاوى وتاريخ المرض الحالي")
        ta("Chief Complaints (C/O)", "الشكاوى الرئيسية", "complaints", 120)
        ti("Onset / Since when?", "متى بدأت الأعراض؟", "onset")
        ta("History of Presenting Illness (HPI)", "تاريخ المرض الحالي", "hpi", 250)
        nav(3, 5)

    elif step == 5:
        sec("Past History", "التاريخ المرضي السابق")
        ta("Previous illnesses / hospitalizations", "الأمراض السابقة / دخول المستشفى", "past_history", 100)
        sec("Family History", "التاريخ العائلي")
        ta("Psychiatric or neurological illness in family", "أمراض نفسية أو عصبية في الأسرة", "family_history", 100)
        nav(4, 6)

    elif step == 6:
        sec("Investigations", "الفحوصات")
        ta("CT, MRI, EEG, IQ tests (SB5, CARS#), etc.", "أشعة مقطعية، رنين، رسم مخ، اختبارات ذكاء (SB5, CARS)...", "investigations", 100)
        sec("Operations and Surgeries", "العمليات والجراحات")
        ta("Previous surgeries", "العمليات الجراحية السابقة", "surgeries", 80)
        sec("Extra Clinical / إضافي")
        rb("Sleep pattern / نمط النوم", "نمط النوم", ["Normal/طبيعي", "Insomnia/أرق", "Hypersomnia/نوم زيادة", "Disrupted/متقطع"], "sleep")
        rb("Appetite / الشهية", "الشهية", ["Normal/طبيعي", "Decreased/قلت", "Increased/زادت"], "appetite")
        ti("Daily screen time / وقت الشاشة اليومي", "وقت الشاشة اليومي", "screen_time")
        ta("Current therapy sessions (speech, skills, etc.)", "الجلسات الحالية (تخاطب، تنمية مهارات...)", "therapy_sessions", 60)
        ta("Additional notes / ملاحظات إضافية", "ملاحظات إضافية", "extra_notes", 80)
        nav(5, 7)

    elif step == 7:
        sec("Child Clinical Checklist", "قائمة التدقيق السريري للأطفال")
        st.caption("Answer Yes / No for each item and add notes where relevant.")

        checklist_items = [
            ("Consanguinity between parents", "القرابة بين الأب والأم"),
            ("Was the child wanted / planned?", "هل الطفل كان مرغوباً فيه؟"),
            ("Was the child's gender (M/F) desired by parents?", "هل نوع الطفل كان مرغوباً فيه؟"),
            ("Motor & cognitive developmental history", "تاريخ النمو الحركي والمعرفي"),
            ("Toilet training age & punishment methods", "سن تدريب دورة المياه وطرق العقاب"),
            ("Siblings at same/different school? Sibling relationship?", "الأخوة في نفس المدرسة؟ علاقتهم ببعض؟"),
            ("Full prenatal / natal / postnatal history", "تاريخ الحمل كامل: قبل/أثناء/بعد الولادة"),
            ("Birth type (natural/CS), forceps/vacuum, incubator, jaundice", "نوع الولادة، جفت/شفاط، حضانة، صفراء"),
            ("Problems during pregnancy / late pregnancy age", "مشاكل أثناء الحمل / سن متأخر للحمل"),
            ("Family members with psychiatric illness, MR, epilepsy", "أقارب لديهم مشكلة نفسية أو إعاقة أو صرع"),
            ("Reaction to stress / punishment methods (especially physical)", "رد الفعل تجاه الضغوط / طرق العقاب"),
            ("If seizures: document doctors and treatments", "في حالة تشنجات: الأطباء والعلاجات"),
            ("High fever (≥40°C) / hospitalization for fever", "ارتفاع حرارة ≥40 درجة / دخول مستشفى حميات"),
            ("Head trauma (location, vomiting, excess sleep/no sleep)", "ارتطام الرأس: مكانه، قيء، نوم زيادة أو عدم نوم"),
            ("Convulsions / post-vaccine complications (esp. MMR at 18m)", "تشنجات / مضاعفات بعد التطعيم (MMR عند سنة ونصف)"),
            ("Cognitive ability distinction: attention vs concentration", "التفرقة بين القدرات المعرفية: انتباه، تركيز، إدراك، فهم"),
            ("Current therapy sessions (speech, skills development)", "جلسات تخاطب / تنمية مهارات"),
            ("Death of a sibling: details, age at time, child's reaction", "وفاة أحد الأخوة: التفاصيل، عمر الطفل، رد فعله"),
            ("Investigations: who ordered them / who reviewed them?", "الفحوصات: من طلبها؟ من شافها؟ (CT, MRI, SB5, CARS)"),
        ]

        checklist_results = {}
        for en, ar in checklist_items:
            st.markdown(f"**{en}**")
            st.markdown(f'<div class="ar">{ar}</div>', unsafe_allow_html=True)
            c1, c2 = st.columns([1, 3])
            with c1:
                ans = st.radio("", ["Yes/نعم", "No/لا", "N/A"], key=f"chk_{en}", horizontal=True, label_visibility="collapsed")
            with c2:
                note = st.text_input("Notes / ملاحظات", key=f"chk_note_{en}", label_visibility="collapsed", placeholder="Notes / ملاحظات")
            checklist_results[en] = {"ar": ar, "answer": ans, "notes": note}
            st.divider()

        st.session_state["checklist"] = checklist_results
        nav(6, final=True)

# ════════════════════════════════════════════════════════
#  GENERATE REPORT
# ════════════════════════════════════════════════════════
if st.session_state.get("generate"):
    st.session_state.generate = False
    s = st.session_state

    def sv(key, default="—"):
        v = s.get(key, "")
        return v.strip() if v and v.strip() else default

    siblings = s.get("siblings", [])
    sibling_text = "\n".join([
        f"  {i+1}. {sib['name']} | {sib['gender']} | Age: {sib['age']} | {sib['edu']} | Notes: {sib['notes']}"
        for i, sib in enumerate(siblings) if sib.get('name')
    ]) or "—"

    if sheet_type == "adult":
        data_block = f"""
PATIENT: {sv('name')} | Age: {sv('age')} | Gender: {sv('gender')}
Date: {sv('taken_date')} | History by: {history_by or '—'} | History Type: {sv('history_type')}
Phone: {sv('phone')} | Referral: {sv('referral')}
Occupation: {sv('occupation')} | Education: {sv('education')}
Social Status: {sv('social_status')} | Hobbies: {sv('hobbies')} | Smoking: {sv('smoking')}

FAMILY:
Father: {sv('father_name')}, Age {sv('father_age')}, Occupation: {sv('father_occ')}
Mother: {sv('mother_name')}, Age {sv('mother_age')}, Occupation: {sv('mother_occ')}
Consanguinity: {sv('consanguinity')} | Chronic illness: {sv('chronic_illness')}

MARRIAGE:
Spouse: {sv('spouse_name')}, Age {sv('spouse_age')}, Occupation: {sv('spouse_occ')}
Duration: {sv('marriage_duration')} | Engagement: {sv('engagement')}
Katb Ketab: {sv('katb_ketab')} | Pre-marriage relation: {sv('pre_marriage_rel')}
Number of children: {sv('num_children')}

SIBLINGS:
{sibling_text}

COMPLAINTS (C/O): {sv('complaints')}
ONSET: {sv('onset')}
HPI: {sv('hpi')}

DRUG HISTORY: {sv('drug_history')}
PAST HISTORY: {sv('past_history')}
FAMILY HISTORY (psychiatric): {sv('family_history')}

INVESTIGATIONS: {sv('investigations')}
SURGERIES: {sv('surgeries')}

SLEEP: {sv('sleep')} | APPETITE: {sv('appetite')}
SUICIDAL IDEATION: {sv('suicidal')}
SUBSTANCE USE: {sv('substance')} — {sv('substance_details')}
EXTRA NOTES: {sv('extra_notes')}
"""
    else:
        checklist = s.get("checklist", {})
        chk_text = "\n".join([
            f"  - {en} ({v['ar']}): {v['answer']} | Notes: {v['notes'] or '—'}"
            for en, v in checklist.items()
        ])
        data_block = f"""
PATIENT: {sv('name')} | Age: {sv('age')} | Gender: {sv('gender')}
Date: {sv('taken_date')} | History by: {history_by or '—'}
Phone: {sv('phone')} | Lives with: {sv('lives_with')}
School: {sv('school_name')} | Grade: {sv('grade')} | Academic performance: {sv('academic')}
Screen time: {sv('screen_time')}

DEVELOPMENTAL MILESTONES:
Pregnancy: {sv('pregnancy')} | Birth: {sv('birth')} | Breastfeeding: {sv('breastfeeding')}
Weaning: {sv('weaning')} | Motor: {sv('motor_dev')} | Teething: {sv('teething')}
Speech: {sv('speech')} | Toilet training: {sv('toilet_training')}
Vaccinations: {sv('vaccinations')} | Vaccination notes: {sv('vaccination_notes')}
Developmental notes: {sv('dev_notes')}

FAMILY:
Father: {sv('father_name')}, Age {sv('father_age')}, Occupation: {sv('father_occ')}, Hereditary illness: {sv('father_hereditary')}
Mother: {sv('mother_name')}, Age {sv('mother_age')}, Occupation: {sv('mother_occ')}, Hereditary illness: {sv('mother_hereditary')}
Consanguinity: {sv('consanguinity')} | Parents relationship: {sv('parents_relation')}

SIBLINGS:
{sibling_text}

COMPLAINTS (C/O): {sv('complaints')}
ONSET: {sv('onset')}
HPI: {sv('hpi')}

PAST HISTORY: {sv('past_history')}
FAMILY HISTORY (psychiatric): {sv('family_history')}

INVESTIGATIONS: {sv('investigations')}
SURGERIES: {sv('surgeries')}
THERAPY SESSIONS: {sv('therapy_sessions')}

SLEEP: {sv('sleep')} | APPETITE: {sv('appetite')}
EXTRA NOTES: {sv('extra_notes')}

CHILD CHECKLIST:
{chk_text}
"""

    patient_name = sv('name', 'Patient')

    prompt = f"""You are a senior consultant psychiatrist. Based on the structured history data below, generate a comprehensive bilingual (Arabic and English) psychiatric history report.

The report must have TWO clearly separated parts:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PART 1 — PROFESSIONAL SUMMARY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Write a professional clinical summary in BOTH Arabic and English.
Structure it with these sections (write each section title in both languages):
- Patient Overview / نظرة عامة عن المريض
- Chief Complaint & Presenting Illness / الشكوى الرئيسية وتاريخ المرض الحالي
- Personal & Social Background / الخلفية الشخصية والاجتماعية
- Family Background / الخلفية العائلية
- Medical & Drug History / التاريخ الطبي والدوائي
- Clinical Observations / الملاحظات السريرية
- Summary Impression / الانطباع العام

Be professional, concise, and clinically accurate. Use formal Arabic and English.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PART 2 — DETAILED RECORD (exact wording)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Present ALL the data as a clean structured table with three columns:
| Field (English) | Field (Arabic) | Client's Response |

Keep the client's exact wording. If they wrote in Arabic, keep it in Arabic. If English, keep in English.
Include every field from the data, including the checklist if present.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
HISTORY DATA:
{data_block}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

History taken by: {history_by or 'Psychologist'}
Sheet type: {sheet_type.upper()}
"""

    with st.spinner("Generating report... / جاري إنشاء التقرير..."):
        try:
            client = Groq(api_key=groq_key)
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2048
            )
            report_text = response.choices[0].message.content
            st.session_state.report_text = report_text
            st.session_state.patient_name_final = patient_name
        except Exception as e:
            st.error(f"Report generation error: {str(e)}")

# ════════════════════════════════════════════════════════
#  BUILD AND DISPLAY REPORT
# ════════════════════════════════════════════════════════
if st.session_state.get("report_text"):
    report_text = st.session_state.report_text
    patient_name = st.session_state.get("patient_name_final", "Patient")

    st.divider()
    st.markdown("### ✅ Report Generated / تم إنشاء التقرير")
    st.text_area("", value=report_text, height=500, label_visibility="collapsed")

    # ── BUILD DOCX ──
    def build_docx(report_text, patient_name, sheet_type, history_by, logo_path, doctor):
        doc = Document()

        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)
            section.different_first_page_header_footer = True
            for hdr in [section.header, section.first_page_header]:
                for p in hdr.paragraphs:
                    p.clear()

        # Page border
        def add_border(doc, color="1B2A4A", size=12):
            for section in doc.sections:
                sectPr = section._sectPr
                pgBorders = OxmlElement('w:pgBorders')
                pgBorders.set(qn('w:offsetFrom'), 'page')
                for side in ('top', 'left', 'bottom', 'right'):
                    b = OxmlElement(f'w:{side}')
                    b.set(qn('w:val'), 'single')
                    b.set(qn('w:sz'), str(size))
                    b.set(qn('w:space'), '24')
                    b.set(qn('w:color'), color)
                    pgBorders.append(b)
                sectPr.append(pgBorders)
        add_border(doc)

        # Page numbers
        def add_page_numbers(doc):
            for section in doc.sections:
                footer = section.footer
                para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                para.clear()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.font.size = Pt(9)
                run.font.color.rgb = CLINIC_BLUE
                for tag, text in [('begin', None), (None, ' PAGE '), ('end', None)]:
                    if tag:
                        el = OxmlElement('w:fldChar')
                        el.set(qn('w:fldCharType'), tag)
                        run._r.append(el)
                    else:
                        instr = OxmlElement('w:instrText')
                        instr.text = text
                        run._r.append(instr)
        add_page_numbers(doc)

        # First page: logo + title
        p_top = doc.add_paragraph()
        p_top.paragraph_format.space_before = Pt(0)
        p_top.paragraph_format.space_after = Pt(6)
        if os.path.exists(logo_path):
            p_top.add_run().add_picture(logo_path, width=Inches(1.2))
        r_title = p_top.add_run(f"   Clinical History Report")
        r_title.font.name = "Arial"; r_title.font.size = Pt(20)
        r_title.font.bold = True; r_title.font.color.rgb = CLINIC_BLUE
        pPr = p_top._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
        bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '1A5CB8')
        pBdr.append(bot); pPr.append(pBdr)

        # Patient info line
        doc.add_paragraph()
        p_info = doc.add_paragraph()
        for label, val in [("Patient: ", patient_name), ("   |   Sheet: ", sheet_type.capitalize()),
                            ("   |   History by: ", history_by or "—")]:
            r = p_info.add_run(label); r.bold = True
            r.font.size = Pt(11); r.font.name = "Arial"; r.font.color.rgb = CLINIC_BLUE
            r2 = p_info.add_run(val); r2.font.size = Pt(11); r2.font.name = "Arial"
        doc.add_paragraph()

        # Report body
        in_table_section = False
        table = None
        for line in report_text.split('\n'):
            line_stripped = line.strip()
            if not line_stripped:
                if not in_table_section:
                    doc.add_paragraph()
                continue

            # Detect table rows
            if line_stripped.startswith('|') and line_stripped.endswith('|'):
                cells = [c.strip() for c in line_stripped.strip('|').split('|')]
                if all(set(c) <= set('-: ') for c in cells):
                    continue  # separator row
                if not in_table_section:
                    in_table_section = True
                    table = doc.add_table(rows=0, cols=3)
                    table.style = 'Table Grid'
                    table.autofit = True
                row = table.add_row()
                for i, cell_text in enumerate(cells[:3]):
                    cell = row.cells[i]
                    cell.text = cell_text
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.name = "Arial"
                continue
            else:
                in_table_section = False
                table = None

            # Section dividers
            if line_stripped.startswith('━'):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                b = OxmlElement('w:bottom')
                b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '1'); b.set(qn('w:color'), '1A5CB8')
                pBdr.append(b); pPr.append(pBdr)
                continue

            # PART headings
            if line_stripped.startswith('PART ') or line_stripped.startswith('━'):
                p = doc.add_paragraph()
                r = p.add_run(line_stripped)
                r.bold = True; r.font.size = Pt(13)
                r.font.name = "Arial"; r.font.color.rgb = CLINIC_BLUE
                continue

            # Section headings (contain / or — or are ALL CAPS short)
            if ('/' in line_stripped and len(line_stripped) < 80 and not line_stripped.startswith('-')) or \
               (line_stripped.isupper() and len(line_stripped) < 60):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                r = p.add_run(line_stripped)
                r.bold = True; r.font.size = Pt(11)
                r.font.name = "Arial"; r.font.color.rgb = CLINIC_BLUE
                pPr2 = p._p.get_or_add_pPr()
                pBdr2 = OxmlElement('w:pBdr')
                b2 = OxmlElement('w:bottom')
                b2.set(qn('w:val'), 'single'); b2.set(qn('w:sz'), '4')
                b2.set(qn('w:space'), '1'); b2.set(qn('w:color'), '1A5CB8')
                pBdr2.append(b2); pPr2.append(pBdr2)
                continue

            # Bullet
            if line_stripped.startswith('- ') or line_stripped.startswith('• '):
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run(line_stripped.lstrip('-•').strip())
                r.font.size = Pt(11); r.font.name = "Arial"
                continue

            # Normal paragraph
            p = doc.add_paragraph()
            r = p.add_run(line_stripped)
            r.font.size = Pt(11); r.font.name = "Arial"

        # Doctor footer
        doc.add_paragraph(); doc.add_paragraph()
        p_sep = doc.add_paragraph()
        pPr_sep = p_sep._p.get_or_add_pPr()
        pBdr_sep = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '1'); top.set(qn('w:color'), '1A5CB8')
        pBdr_sep.append(top); pPr_sep.append(pBdr_sep)

        p_dr = doc.add_paragraph()
        r_dr = p_dr.add_run(doctor["name"])
        r_dr.bold = True; r_dr.font.size = Pt(12)
        r_dr.font.name = "Arial"; r_dr.font.color.rgb = CLINIC_BLUE

        for t in ["title1", "title2", "title3", "title4"]:
            p_t = doc.add_paragraph()
            r_t = p_t.add_run(doctor[t])
            r_t.font.size = Pt(10); r_t.font.name = "Arial"
            r_t.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p_t.paragraph_format.space_before = Pt(0)
            p_t.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()
        p_addr = doc.add_paragraph()
        p_addr.add_run(f"📍  {doctor['address']}").font.size = Pt(10)
        p_phone = doc.add_paragraph()
        r_ph = p_phone.add_run(f"📞  {doctor['phone']}")
        r_ph.font.size = Pt(10); r_ph.bold = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    docx_buf = build_docx(report_text, patient_name, sheet_type, history_by, LOGO_PATH, DOCTOR)
    filename = f"{patient_name.replace(' ', '_')}_HistorySheet.docx"

    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            label="📄 Download .docx",
            data=docx_buf,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col2:
        if st.button("📧 Send to Email"):
            if not gmail_user or not gmail_pass:
                st.error("Enter Gmail credentials in the sidebar first.")
            else:
                try:
                    docx_buf2 = build_docx(report_text, patient_name, sheet_type, history_by, LOGO_PATH, DOCTOR)
                    msg = MIMEMultipart()
                    msg['From'] = gmail_user
                    msg['To'] = RECIPIENT_EMAIL
                    msg['Subject'] = f"History Report — {patient_name}"
                    msg.attach(MIMEText(f"Please find attached the history report for {patient_name}.\n\nHistory by: {history_by or '—'}\nSheet type: {sheet_type.capitalize()}", 'plain'))
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(docx_buf2.read())
                    encoders.encode_base64(part)
                    part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
                    msg.attach(part)
                    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                        server.login(gmail_user, gmail_pass)
                        server.sendmail(gmail_user, RECIPIENT_EMAIL, msg.as_string())
                    st.success(f"✅ Report sent to {RECIPIENT_EMAIL}")
                except Exception as e:
                    st.error(f"Email error: {str(e)}")

    with col3:
        if st.button("↺ New Patient / مريض جديد"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
