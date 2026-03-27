import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os, smtplib, re, zipfile
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date

# ══════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════
RECIPIENT_EMAIL = "yusuf.a.abdelatti@gmail.com"
GMAIL_USER      = "yusuf.a.abdelatti@gmail.com"
GMAIL_PASS      = "erjl ehlj wpyg mfgx"
LOGO_PATH       = os.path.join(os.path.dirname(__file__), "logo.png")

# Logo-extracted palette
LOGO_BLUE   = "#A1B6F3"   # periwinkle from logo
DEEP_BLUE   = "#3B5FC0"   # darker variant for text/headings
DARK_BLUE   = "#1B2A4A"   # near-navy for borders
MID_BLUE    = "#6B8ED6"   # mid-tone accent
LIGHT_BG    = "#F0F4FF"   # very light blue tint background
PILL_BG     = "#E8EEFF"   # unselected pill
PILL_ACTIVE = "#3B5FC0"   # selected pill

CLINIC_BLUE_RGB  = RGBColor(0x3B, 0x5F, 0xC0)
DARK_BLUE_RGB    = RGBColor(0x1B, 0x2A, 0x4A)
LOGO_BLUE_RGB    = RGBColor(0xA1, 0xB6, 0xF3)

DOCTOR = {
    "name":   "Dr. Hany Elhennawy",
    "title1": "Consultant of Neuro-Psychiatry",
    "title2": "Aviation Medical Council — Faculty of Medicine, 6th October University",
    "title3": "MD of Neuroscience Research, Karolinska Institute — Sweden",
    "phone":  "+20 1000756200",
}

# ══════════════════════════════════════════════════════════════
#  PAGE CONFIG & CSS
# ══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="أخذ التاريخ المرضي — د. هاني الحناوي",
    page_icon="🧠",
    layout="wide"
)

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@300;400;600;700&family=Inter:wght@300;400;500;600&display=swap');

html, body, [class*="css"] {{
    font-family: 'Inter', 'Cairo', sans-serif;
    background-color: {LIGHT_BG};
}}
.stApp {{ background-color: {LIGHT_BG}; }}

/* ── Header ── */
.clinic-header {{
    background: linear-gradient(135deg, {DEEP_BLUE} 0%, {MID_BLUE} 100%);
    border-radius: 16px;
    padding: 24px 32px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    gap: 20px;
    box-shadow: 0 4px 20px rgba(59,95,192,0.25);
}}
.clinic-header-text h1 {{
    color: white;
    font-size: 22px;
    font-weight: 700;
    margin: 0 0 4px 0;
    font-family: 'Cairo', sans-serif;
}}
.clinic-header-text p {{
    color: rgba(255,255,255,0.78);
    font-size: 13px;
    margin: 0;
}}

/* ── Section headers ── */
.sec-header {{
    font-size: 14px;
    font-weight: 700;
    color: {DEEP_BLUE};
    margin: 28px 0 12px 0;
    padding: 10px 16px;
    background: white;
    border-radius: 10px;
    border-right: 4px solid {LOGO_BLUE};
    border-left: 4px solid {LOGO_BLUE};
    box-shadow: 0 2px 8px rgba(59,95,192,0.08);
    display: flex;
    align-items: center;
    gap: 8px;
    direction: rtl;
}}

/* ── Field labels ── */
.field-label {{
    font-size: 12.5px;
    font-weight: 600;
    color: {DARK_BLUE};
    margin-bottom: 6px;
    direction: rtl;
    text-align: right;
}}

/* ── Pill buttons ── */
.pill-container {{
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
    margin-bottom: 4px;
    direction: rtl;
}}
.pill-btn {{
    display: inline-block;
    padding: 7px 18px;
    border-radius: 50px;
    font-size: 13px;
    font-family: 'Cairo', 'Inter', sans-serif;
    font-weight: 500;
    cursor: pointer;
    border: 2px solid {LOGO_BLUE};
    background: white;
    color: {DEEP_BLUE};
    transition: all 0.15s ease;
    user-select: none;
    white-space: nowrap;
}}
.pill-btn:hover {{
    background: {PILL_BG};
    border-color: {DEEP_BLUE};
}}
.pill-btn.active {{
    background: {PILL_ACTIVE};
    border-color: {PILL_ACTIVE};
    color: white;
    box-shadow: 0 2px 8px rgba(59,95,192,0.35);
}}

/* ── Input fields ── */
div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea {{
    background: white !important;
    border: 1.5px solid #C5D3F5 !important;
    border-radius: 10px !important;
    font-family: 'Cairo', 'Inter', sans-serif !important;
    font-size: 13.5px !important;
    direction: rtl !important;
    text-align: right !important;
    color: {DARK_BLUE} !important;
    transition: border-color 0.2s !important;
}}
div[data-testid="stTextInput"] input:focus,
div[data-testid="stTextArea"] textarea:focus {{
    border-color: {DEEP_BLUE} !important;
    box-shadow: 0 0 0 3px rgba(59,95,192,0.12) !important;
}}

/* ── Sheet type toggle ── */
div[data-testid="stRadio"] > div {{
    gap: 12px !important;
    flex-direction: row !important;
}}
div[data-testid="stRadio"] > div > label {{
    background: white !important;
    border: 2px solid {LOGO_BLUE} !important;
    border-radius: 50px !important;
    padding: 8px 24px !important;
    font-size: 14px !important;
    font-weight: 600 !important;
    font-family: 'Cairo', sans-serif !important;
    color: {DEEP_BLUE} !important;
    cursor: pointer !important;
    transition: all 0.15s !important;
}}
div[data-testid="stRadio"] > div > label:has(input:checked) {{
    background: {DEEP_BLUE} !important;
    color: white !important;
    border-color: {DEEP_BLUE} !important;
}}
div[data-testid="stRadio"] > label {{ display: none; }}

/* ── Buttons ── */
.stButton > button {{
    background: {DEEP_BLUE} !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 10px 28px !important;
    font-family: 'Cairo', 'Inter', sans-serif !important;
    font-size: 14px !important;
    font-weight: 600 !important;
    transition: all 0.2s !important;
    box-shadow: 0 2px 10px rgba(59,95,192,0.25) !important;
}}
.stButton > button:hover {{
    background: {MID_BLUE} !important;
    box-shadow: 0 4px 14px rgba(59,95,192,0.4) !important;
    transform: translateY(-1px) !important;
}}
.stButton > button[kind="primary"] {{
    background: linear-gradient(135deg, {DEEP_BLUE}, {MID_BLUE}) !important;
    font-size: 16px !important;
    padding: 14px 40px !important;
}}

/* ── Cards ── */
.info-card {{
    background: white;
    border-radius: 12px;
    padding: 16px 20px;
    margin-bottom: 16px;
    border: 1px solid #DDE5F8;
    box-shadow: 0 2px 8px rgba(59,95,192,0.06);
}}

/* ── Siblings table header ── */
.sib-col-header {{
    font-size: 11px;
    font-weight: 700;
    color: {MID_BLUE};
    text-align: center;
    letter-spacing: 0.04em;
    text-transform: uppercase;
    padding: 4px 0;
}}

/* ── Access gate ── */
.access-gate {{
    max-width: 420px;
    margin: 4rem auto;
    background: white;
    border-radius: 20px;
    padding: 40px 36px;
    box-shadow: 0 8px 40px rgba(59,95,192,0.15);
    text-align: center;
}}
.access-gate h2 {{
    color: {DEEP_BLUE};
    font-size: 22px;
    font-weight: 700;
    font-family: 'Cairo', sans-serif;
    margin-bottom: 8px;
}}
.access-gate p {{
    color: #666;
    font-size: 13.5px;
    line-height: 1.8;
    margin-bottom: 24px;
    font-family: 'Cairo', sans-serif;
}}

div[data-testid="stDivider"] {{ margin: 20px 0; }}
div[data-testid="stCaption"] {{ color: {MID_BLUE}; font-weight: 500; }}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  ACCESS GATE
# ══════════════════════════════════════════════════════════════
if "access_granted" not in st.session_state:
    st.session_state.access_granted = False

if not st.session_state.access_granted:
    if os.path.exists(LOGO_PATH):
        c1, c2, c3 = st.columns([1,2,1])
        with c2:
            st.image(LOGO_PATH, use_container_width=True)

    st.markdown("""
    <div class="access-gate">
        <h2>استمارة التاريخ المرضي</h2>
        <p>هذه الاستمارة مخصصة للأخصائيين المعتمدين فقط.<br>يرجى إدخال رمز الوصول للمتابعة.</p>
    </div>""", unsafe_allow_html=True)

    col_a, col_b, col_c = st.columns([1,2,1])
    with col_b:
        code = st.text_input("رمز الوصول", type="password",
                             placeholder="أدخل رمز الوصول",
                             label_visibility="collapsed")
        if st.button("دخول ←", use_container_width=True):
            valid_codes = [c.strip() for c in st.secrets.get("ACCESS_CODE","").split(",")]
            if code.strip() in valid_codes:
                st.session_state.access_granted = True
                st.rerun()
            else:
                st.error("⚠ رمز الوصول غير صحيح.")
    st.stop()

# ══════════════════════════════════════════════════════════════
#  PILL BUTTON HELPER
# ══════════════════════════════════════════════════════════════
def pill_select(label_ar: str, label_en: str, options: list, key: str, allow_none: bool = False) -> str:
    """Render pill buttons for single-select. Returns selected value or ''."""
    if key not in st.session_state:
        st.session_state[key] = ""

    current = st.session_state[key]

    # Label
    en_part = f" <span style='color:{MID_BLUE};font-size:11px;font-weight:400;'>{label_en}</span>" if label_en else ""
    st.markdown(f'<div class="field-label">{label_ar}{en_part}</div>', unsafe_allow_html=True)

    # Build pill HTML
    pills_html = '<div class="pill-container">'
    for opt in options:
        active_class = "active" if current == opt else ""
        pills_html += f'<span class="pill-btn {active_class}" id="pill_{key}_{opt}">{opt}</span>'
    pills_html += "</div>"
    st.markdown(pills_html, unsafe_allow_html=True)

    # Use columns as click targets (invisible radio)
    cols = st.columns(len(options))
    for i, opt in enumerate(options):
        with cols[i]:
            if st.button(opt, key=f"_pill_{key}_{i}", help=None,
                         use_container_width=True):
                if allow_none and st.session_state[key] == opt:
                    st.session_state[key] = ""
                else:
                    st.session_state[key] = opt
                st.rerun()

    return st.session_state[key]


def pill_multi(label_ar: str, label_en: str, options: list, key: str) -> list:
    """Multi-select pills. Returns list of selected values."""
    if key not in st.session_state:
        st.session_state[key] = []

    current = list(st.session_state[key])

    en_part = f" <span style='color:{MID_BLUE};font-size:11px;font-weight:400;'>{label_en}</span>" if label_en else ""
    st.markdown(f'<div class="field-label">{label_ar}{en_part}</div>', unsafe_allow_html=True)

    pills_html = '<div class="pill-container">'
    for opt in options:
        active_class = "active" if opt in current else ""
        pills_html += f'<span class="pill-btn {active_class}">{opt}</span>'
    pills_html += "</div>"
    st.markdown(pills_html, unsafe_allow_html=True)

    cols = st.columns(min(len(options), 5))
    for i, opt in enumerate(options):
        with cols[i % len(cols)]:
            if st.button(opt, key=f"_mpill_{key}_{i}", use_container_width=True):
                new = list(st.session_state[key])
                if opt in new:
                    new.remove(opt)
                else:
                    new.append(opt)
                st.session_state[key] = new
                st.rerun()

    return st.session_state[key]


# ══════════════════════════════════════════════════════════════
#  FIELD HELPERS
# ══════════════════════════════════════════════════════════════
def sec(icon: str, ar: str, en: str = ""):
    en_part = f" / {en}" if en else ""
    st.markdown(f'<div class="sec-header">{icon} {ar}{en_part}</div>',
                unsafe_allow_html=True)

def lbl(ar: str, en: str = ""):
    en_part = f" <span style='color:{MID_BLUE};font-size:11px;font-weight:400;'>{en}</span>" if en else ""
    st.markdown(f'<div class="field-label">{ar}{en_part}</div>', unsafe_allow_html=True)

def ti(ar: str, en: str, key: str, placeholder: str = "") -> str:
    lbl(ar, en)
    return st.text_input("", key=key, placeholder=placeholder,
                         label_visibility="collapsed")

def ta(ar: str, en: str, key: str, height: int = 110) -> str:
    lbl(ar, en)
    return st.text_area("", key=key, height=height,
                        label_visibility="collapsed")

def sv(d: dict, key: str, default: str = "لم يُذكر") -> str:
    v = d.get(key, "")
    if not v:
        return default
    if isinstance(v, list):
        return "، ".join(v) if v else default
    v = str(v).strip()
    return v if v and v not in ["—", "— اختر —", "لم يُذكر", ""] else default


# ══════════════════════════════════════════════════════════════
#  CHOICE LISTS
# ══════════════════════════════════════════════════════════════
GENDER_AR    = ["ذكر", "أنثى"]
EDU_AR       = ["أمي","ابتدائي","إعدادي","ثانوي","جامعي","دراسات عليا"]
OCC_AR       = ["موظف حكومي","موظف قطاع خاص","أعمال حرة","طالب","ربة منزل","متقاعد","عاطل","أخرى"]
SOCIAL_AR    = ["أعزب","متزوج","مطلق","أرمل","منفصل"]
SMOKING_AR   = ["لا يدخن","مدخن","توقف","شيشة","مدخن + شيشة"]
REFERRAL_AR  = ["ذاتي","الأسرة","طبيب","أخصائي نفسي","مدرسة","أخرى"]
HTYPE_AR     = ["أولي","متابعة","طارئ","استشاري"]
ALIVE_M      = ["على قيد الحياة","متوفى","غير معروف"]
ALIVE_F      = ["على قيد الحياة","متوفاة","غير معروف"]
CONS_AR      = ["لا توجد قرابة","درجة أولى","درجة ثانية","درجة ثالثة"]
PARENTS_REL  = ["جيدة","متوسطة","سيئة","منفصلان","مطلقان","أحدهما متوفى"]
MARQ_AR      = ["جيدة","متوسطة","سيئة","منفصلان"]
PRE_MAR      = ["لا توجد علاقة سابقة","تعارف فقط","علاقة طويلة","زواج مرتب","أخرى"]
NUM_CHILD    = ["لا يوجد","1","2","3","4","5","6 فأكثر"]
MARRIAGE_DUR = ["أقل من سنة","1-3 سنوات","3-5 سنوات","5-10 سنوات","أكثر من 10 سنوات"]
ENGAGEMENT   = ["لم تكن هناك خطوبة","أقل من 3 أشهر","3-6 أشهر","6-12 شهراً","أكثر من سنة"]
ONSET_MODE   = ["مفاجئ","تدريجي"]
COURSE_AR    = ["مستمر","نوبات متكررة","في تحسن","في تدهور","متذبذب"]
COMPLIANCE   = ["ملتزم","غير منتظم","غير ملتزم","رافض"]
INSIGHT_AR   = ["كاملة","جزئية","غائبة"]
SLEEP_AR     = ["طبيعي","أرق","نوم زيادة","متقطع"]
APPETITE_AR  = ["طبيعية","قلت","زادت"]
SUICIDAL_AR  = ["لا توجد","أفكار سلبية فقط","أفكار نشطة","خطة واضحة"]
SUBSTANCE_AR = ["لا يوجد","كحول","حشيش","حبوب مهدئة","متعدد","أخرى"]
HOBBIES_AR   = ["قراءة","رياضة","موسيقى","رسم","طبخ","ألعاب إلكترونية","تواصل اجتماعي","لا توجد","أخرى"]
CHRONIC_AR   = ["لا يوجد","سكري","ضغط","أمراض قلب","أمراض كلى","أمراض مناعية","سرطان","أخرى"]
SIB_GENDER   = ["ذكر","أنثى","—"]
SIB_EDU      = ["روضة","ابتدائي","إعدادي","ثانوي","جامعي","خريج","لا يدرس"]
BIRTH_ORDER  = ["الأول","الثاني","الثالث","الرابع","الخامس","السادس فأكثر","وحيد"]
BIRTH_TYPE   = ["طبيعي","قيصري","بالجفت","بالشفاط"]
BIRTH_COMP   = ["لا يوجد","صفراء","حضانة","اختناق","وزن منخفض","أخرى"]
BF_AR        = ["رضاعة طبيعية","رضاعة صناعية","مختلطة"]
WEANING_AR   = ["قبل 6 أشهر","6-12 شهراً","12-18 شهراً","18-24 شهراً","بعد سنتين"]
MOTOR_AR     = ["طبيعي","متأخر","مبكر"]
SPEECH_AR    = ["طبيعي","متأخر","غائب","تراجع بعد اكتمال"]
TEETH_AR     = ["طبيعي (6-8 أشهر)","مبكر","متأخر (بعد 12 شهراً)"]
TOILET_AR    = ["طبيعي (18-30 شهراً)","مبكر","متأخر (بعد 3 سنوات)"]
VACC_AR      = ["مكتمل","غير مكتمل","غير معروف"]
ACADEMIC_AR  = ["ممتاز","جيد","متوسط","ضعيف","لا يدرس"]
WANTED_AR    = ["نعم، مرغوب فيه","لا، لم يكن مرغوباً","حمل غير مخطط"]
GENDER_DES   = ["كان النوع مرغوباً","كان يُفضَّل نوع آخر","لا فرق"]
LIVES_WITH   = ["مع الوالدين","مع الأم فقط","مع الأب فقط","مع الجدين","مع أقارب","أخرى"]
SCREEN_AR    = ["أقل من ساعة","1-2 ساعة","2-4 ساعات","4-6 ساعات","أكثر من 6 ساعات"]
PREG_AR      = ["طبيعي","مع ضغط","مع سكري","مع نزيف","سن متأخرة (>35)","مشكلة أخرى"]
YES_NO_NA    = ["نعم","لا","لا ينطبق"]
YES_NO       = ["نعم","لا"]
PUNISHMENT   = ["لفظي","حرمان من امتيازات","جسدي","تجاهل","عقاب اجتماعي","أخرى"]
STRESS_RXN   = ["هادئ","بكاء","عدوان","انسحاب","نوبات غضب","تبوّل لاإرادي","أخرى"]
SIB_REL      = ["جيدة","متوسطة","تنافسية","صراع مستمر","إهمال متبادل"]
SAME_SCH     = ["نعم","لا","لا ينطبق"]

# ══════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown(f"""
    <div style="text-align:center;padding:16px 0 8px;">
        <div style="font-size:28px;">🧠</div>
        <div style="font-size:15px;font-weight:700;color:{DEEP_BLUE};font-family:'Cairo',sans-serif;">
            د. هاني الحناوي
        </div>
        <div style="font-size:11px;color:#888;margin-top:4px;">
            Psychiatric Center
        </div>
    </div>
    <hr style="border:none;border-top:1px solid #DDE5F8;margin:12px 0;">
    """, unsafe_allow_html=True)

    history_by = st.text_input("👤 اسم الأخصائي / Psychologist Name",
                               placeholder="الاسم الكامل")

    st.markdown("<hr style='border:none;border-top:1px solid #DDE5F8;'>", unsafe_allow_html=True)

    if st.button("↺ مريض جديد / New Patient", use_container_width=True):
        preserve = {"access_granted", "access_granted"}
        for key in list(st.session_state.keys()):
            if key not in preserve:
                del st.session_state[key]
        st.rerun()

# ══════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════
col_logo, col_title = st.columns([1, 5])
with col_logo:
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, width=130)
with col_title:
    st.markdown(f"""
    <div style="padding:8px 0;">
        <div style="font-size:24px;font-weight:700;color:{DEEP_BLUE};
                    font-family:'Cairo',sans-serif;line-height:1.3;">
            استمارة أخذ التاريخ المرضي
        </div>
        <div style="font-size:13px;color:{MID_BLUE};margin-top:4px;font-family:'Cairo',sans-serif;">
            عيادة د. هاني الحناوي — طب وجراحة الأعصاب والنفس
        </div>
    </div>
    """, unsafe_allow_html=True)

st.divider()

# ══════════════════════════════════════════════════════════════
#  SHEET TYPE
# ══════════════════════════════════════════════════════════════
sheet_type = st.radio(
    "**نوع الاستمارة / Sheet Type**",
    ["👤 بالغ / Adult", "👶 طفل / Child"],
    horizontal=True
)
is_adult = "بالغ" in sheet_type
st.divider()
d = {}

# ══════════════════════════════════════════════════════════════
#  ██████  ADULT FORM  ██████
# ══════════════════════════════════════════════════════════════
if is_adult:

    sec("👤","البيانات الشخصية","Personal Details")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<div class="info-card">', unsafe_allow_html=True)
        d["name"]       = ti("الاسم الكامل","Full Name","a_name")
        d["birthdate"]  = ti("تاريخ الميلاد","Birth Date","a_birthdate", "DD/MM/YYYY")
        _bd = st.session_state.get("a_birthdate","")
        _age_str = ""
        if _bd:
            try:
                import re as _re
                from datetime import date as _date
                _parts = _re.split(r'[/\-\.]', _bd.strip())
                if len(_parts)==3:
                    _d,_m,_y = int(_parts[0]),int(_parts[1]),int(_parts[2])
                    _t = _date.today()
                    _y2 = _t.year - _y - ((_t.month,_t.day) < (_m,_d))
                    _mo = (_t.month - _m) % 12
                    _age_str = f"{_y2} years, {_mo} months"
            except: pass
        d["age"] = _age_str
        if _age_str:
            st.caption(f"🎂 {_age_str}")
        d["gender"]    = pill_select("النوع","Gender", GENDER_AR, "a_gender")
        d["education"] = pill_select("المستوى التعليمي","Education", EDU_AR, "a_edu")
        st.markdown('</div>', unsafe_allow_html=True)

    with c2:
        st.markdown('<div class="info-card">', unsafe_allow_html=True)
        d["occupation"]  = pill_select("الوظيفة","Occupation", OCC_AR, "a_occ")
        d["occ_detail"]  = ti("تفاصيل الوظيفة","Occupation details","a_occd")
        d["social"]      = pill_select("الحالة الاجتماعية","Social Status", SOCIAL_AR, "a_social")
        d["smoking"]     = pill_select("التدخين","Smoking", SMOKING_AR, "a_smoking")
        d["phone"]       = ti("رقم الهاتف","Phone","a_phone")
        d["date"]        = ti("تاريخ الجلسة","Date","a_date", str(date.today()))
        st.markdown('</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        d["referral"] = pill_select("مصدر الإحالة","Referral", REFERRAL_AR, "a_referral")
    with c2:
        d["htype"]    = pill_select("نوع التاريخ","History Type", HTYPE_AR, "a_htype")

    d["hobbies"] = pill_multi("الهوايات","Hobbies", HOBBIES_AR, "a_hobbies")

    sec("👨‍👩‍👧","بيانات الأسرة","Family Details")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin-bottom:8px;">الأب / Father</div>', unsafe_allow_html=True)
        d["father_name"]  = ti("الاسم","Name","a_fn")
        d["father_age"]   = ti("السن","Age","a_fa")
        d["father_occ"]   = ti("الوظيفة","Occupation","a_fo")
        d["father_alive"] = pill_select("الحالة","Status", ALIVE_M, "a_falive")
    with c2:
        st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin-bottom:8px;">الأم / Mother</div>', unsafe_allow_html=True)
        d["mother_name"]  = ti("الاسم","Name","a_mn")
        d["mother_age"]   = ti("السن","Age","a_ma")
        d["mother_occ"]   = ti("الوظيفة","Occupation","a_mo")
        d["mother_alive"] = pill_select("الحالة","Status", ALIVE_F, "a_malive")

    c1, c2 = st.columns(2)
    with c1:
        d["consanguinity"]    = pill_select("القرابة بين الأبوين","Consanguinity", CONS_AR, "a_cons")
    with c2:
        d["parents_together"] = pill_select("هل الأبوان يعيشان معاً؟","Living together?", YES_NO_NA, "a_ptog")
    d["chronic"] = pill_select("مرض مزمن في الأسرة","Chronic illness", CHRONIC_AR, "a_chronic")

    sec("💍","بيانات الزواج","Marriage Details")
    c1, c2 = st.columns(2)
    with c1:
        d["spouse_name"]  = ti("اسم الزوج/الزوجة","Spouse Name","a_spn")
        d["spouse_age"]   = ti("السن","Age","a_spa")
        d["spouse_occ"]   = pill_select("الوظيفة","Occupation", OCC_AR, "a_spo")
        d["marriage_dur"] = pill_select("مدة الزواج","Duration", MARRIAGE_DUR, "a_mdur")
    with c2:
        d["engagement"]   = pill_select("فترة الخطوبة","Engagement", ENGAGEMENT, "a_eng")
        d["num_children"] = pill_select("عدد الأبناء","Children", NUM_CHILD, "a_nch")
        d["katb"]         = pill_select("كتب كتاب قبل الزواج؟","Katb Ketab?", YES_NO_NA, "a_katb")
        d["marriage_qual"]= pill_select("جودة الزواج","Quality", MARQ_AR, "a_mqual")
        d["pre_marriage"] = pill_select("العلاقة قبل الزواج","Before marriage", PRE_MAR, "a_pre")

    sec("👫","الإخوة والأخوات","Brothers & Sisters")
    siblings = []
    sib_cols = st.columns([1,2,1,2,2])
    headers = ["النوع","الاسم","السن","التعليم","ملاحظات"]
    for col, h in zip(sib_cols, headers):
        col.markdown(f'<div class="sib-col-header">{h}</div>', unsafe_allow_html=True)

    for i in range(1, 6):
        c1,c2,c3,c4,c5 = st.columns([1,2,1,2,2])
        with c1:
            g  = st.selectbox("",SIB_GENDER,key=f"a_sg{i}",label_visibility="collapsed")
        with c2:
            n  = st.text_input("",key=f"a_sn{i}",placeholder=f"الاسم {i}",label_visibility="collapsed")
        with c3:
            ag = st.text_input("",key=f"a_sa{i}",placeholder="",label_visibility="collapsed")
        with c4:
            e  = st.selectbox("",["—"]+SIB_EDU,key=f"a_se{i}",label_visibility="collapsed")
        with c5:
            nt = st.text_input("",key=f"a_st{i}",placeholder="",label_visibility="collapsed")
        if n.strip():
            siblings.append({"gender":g,"name":n,"age":ag,"edu":e,"notes":nt})
    d["siblings"] = siblings

    sec("🩺","الشكاوى وتاريخ المرض الحالي","Complaints & HPI")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["onset"]      = ti("متى بدأت الأعراض؟","Onset","a_onset")
    with c2:
        d["onset_mode"] = pill_select("طريقة البداية","Mode", ONSET_MODE, "a_omode")
    with c3:
        d["course"]     = pill_select("مسار المرض","Course", COURSE_AR, "a_course")

    d["complaints"] = ta("الشكاوى الرئيسية (C/O)","Chief Complaints","a_co",120)
    d["hpi"]        = ta("تاريخ المرض الحالي بالتفصيل (HPI)","History of Presenting Illness","a_hpi",220)

    sec("💊","تاريخ الأدوية","Drug History")
    c1, c2 = st.columns(2)
    with c1:
        d["on_meds"]    = pill_select("هل يتناول أدوية حالياً؟","On medication?", YES_NO_NA, "a_onmeds")
    with c2:
        d["compliance"] = pill_select("الالتزام","Compliance", COMPLIANCE, "a_comp")
    d["drug_hx"]        = ta("تفاصيل الأدوية (الاسم، الجرعة، المدة)","Medications","a_drug",100)

    sec("📋","التاريخ المرضي السابق","Past History")
    c1, c2 = st.columns(2)
    with c1:
        d["prev_psych"] = pill_select("مرض نفسي سابق؟","Previous psychiatric?", YES_NO_NA, "a_ppsych")
    with c2:
        d["prev_hosp"]  = pill_select("دخول مستشفى سابق؟","Previous hospitalization?", YES_NO_NA, "a_phosp")
    d["past_hx"] = ta("تفاصيل التاريخ السابق","Details","a_past",80)

    sec("🧬","التاريخ العائلي","Family History")
    c1, c2 = st.columns(2)
    with c1:
        d["fam_psych"] = pill_select("مرض نفسي في الأسرة؟","Psychiatric in family?", YES_NO_NA, "a_fpsych")
        if st.session_state.get("a_fpsych") == "نعم":
            d["fam_psych_details"] = ti("تفاصيل المرض النفسي","Details","a_fpsych_det")
        else:
            d["fam_psych_details"] = ""
    with c2:
        d["fam_neuro"] = pill_select("مرض عصبي في الأسرة؟","Neurological?", YES_NO_NA, "a_fneuro")
        if st.session_state.get("a_fneuro") == "نعم":
            d["fam_neuro_details"] = ti("تفاصيل المرض العصبي","Details","a_fneuro_det")
        else:
            d["fam_neuro_details"] = ""
    d["family_hx"] = ta("تفاصيل التاريخ العائلي","Details","a_famhx",80)

    sec("🔬","الفحوصات","Investigations")
    d["had_inv"]        = pill_select("هل أُجريت فحوصات؟","Done?", YES_NO_NA, "a_hadinv")
    d["investigations"] = ta("تفاصيل الفحوصات ونتائجها","Details","a_inv",80)

    sec("🏥","العمليات والجراحات","Operations & Surgeries")
    d["had_surg"]  = pill_select("عمليات جراحية سابقة؟","Previous surgeries?", YES_NO_NA, "a_hsurg")
    d["surgeries"] = ta("تفاصيل العمليات","Details","a_surg",60)

    sec("📊","التقييم السريري","Clinical Assessment")
    c1, c2 = st.columns(2)
    with c1:
        d["sleep"]     = pill_select("نمط النوم","Sleep", SLEEP_AR, "a_sleep")
        d["appetite"]  = pill_select("الشهية","Appetite", APPETITE_AR, "a_appetite")
        d["suicidal"]  = pill_select("أفكار انتحارية","Suicidal ideation", SUICIDAL_AR, "a_suicidal")
        d["insight"]   = pill_select("البصيرة","Insight", INSIGHT_AR, "a_insight")
    with c2:
        d["substance"]         = pill_select("تعاطي مواد","Substance use", SUBSTANCE_AR, "a_subs")
        d["substance_details"] = ta("تفاصيل المواد","Details","a_subsd",60)
    d["extra_notes"] = ta("ملاحظات إضافية","Additional notes","a_extra",80)
    patient_name = d.get("name") or "المريض"


# ══════════════════════════════════════════════════════════════
#  ██████  CHILD FORM  ██████
# ══════════════════════════════════════════════════════════════
else:

    sec("👶","البيانات الشخصية","Personal Details")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<div class="info-card">', unsafe_allow_html=True)
        d["name"]        = ti("اسم الطفل كاملاً","Child's Full Name","c_name")
        d["birthdate"]   = ti("تاريخ الميلاد","Birth Date","c_birthdate","DD/MM/YYYY")
        _bd2 = st.session_state.get("c_birthdate","")
        _age_str2 = ""
        if _bd2:
            try:
                import re as _re2
                from datetime import date as _date2
                _p2 = _re2.split(r'[/\-\.]', _bd2.strip())
                if len(_p2)==3:
                    _d2,_m2,_y2 = int(_p2[0]),int(_p2[1]),int(_p2[2])
                    _t2 = _date2.today()
                    _yr2 = _t2.year - _y2 - ((_t2.month,_t2.day) < (_m2,_d2))
                    _mo2 = (_t2.month - _m2) % 12
                    _age_str2 = f"{_yr2} years, {_mo2} months"
            except: pass
        d["age"] = _age_str2
        if _age_str2:
            st.caption(f"🎂 {_age_str2}")
        d["gender"]      = pill_select("النوع","Gender", GENDER_AR, "c_gender")
        d["birth_order"] = pill_select("ترتيب الميلاد","Birth Order", BIRTH_ORDER, "c_border")
        d["lives_with"]  = pill_select("يعيش مع","Lives with", LIVES_WITH, "c_lives")
        st.markdown('</div>', unsafe_allow_html=True)

    with c2:
        st.markdown('<div class="info-card">', unsafe_allow_html=True)
        d["school"]      = ti("اسم المدرسة","School Name","c_school")
        d["grade"]       = ti("الصف الدراسي","Grade","c_grade")
        d["academic"]    = pill_select("المستوى الدراسي","Academic Level", ACADEMIC_AR, "c_academic")
        d["screen_time"] = pill_select("وقت الشاشة اليومي","Screen Time", SCREEN_AR, "c_screen")
        d["phone"]       = ti("تليفون","Phone","c_phone")
        d["date"]        = ti("تاريخ الجلسة","Date","c_date", str(date.today()))
        st.markdown('</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        d["wanted"]     = pill_select("هل كان الطفل مرغوباً فيه؟","Wanted?", WANTED_AR, "c_wanted")
    with c2:
        d["gender_des"] = pill_select("هل النوع كان مرغوباً؟","Gender desired?", GENDER_DES, "c_gdes")
    c1, c2 = st.columns(2)
    with c1:
        d["referral"]   = pill_select("مصدر الإحالة","Referral", REFERRAL_AR, "c_referral")
    with c2:
        d["htype"]      = pill_select("نوع التاريخ","History Type", HTYPE_AR, "c_htype")

    sec("🌱","مراحل النمو","Developmental History")
    st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin:12px 0 8px;">الحمل والولادة / Prenatal & Natal</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        d["pregnancy"]   = pill_select("الحمل","Pregnancy", PREG_AR, "c_preg")
        d["birth_type"]  = pill_select("نوع الولادة","Birth type", BIRTH_TYPE, "c_btype")
    with c2:
        d["birth_comp"]  = pill_select("مضاعفات الولادة","Birth complications", BIRTH_COMP, "c_bcomp")
        d["vacc_status"] = pill_select("التطعيمات","Vaccinations", VACC_AR, "c_vacc")
    with c3:
        d["vacc_comp"]   = ti("مضاعفات بعد التطعيم","Post-vaccine comp.","c_vcomp")
    d["preg_notes"] = ta("ملاحظات الحمل والولادة","Prenatal/natal notes","c_pregnotes",70)

    st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin:16px 0 8px;">التغذية والنمو / Feeding & Growth</div>', unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        d["breastfeeding"] = pill_select("الرضاعة","Breastfeeding", BF_AR, "c_bf")
    with c2:
        d["weaning"]       = pill_select("سن الفطام","Weaning", WEANING_AR, "c_wean")
    with c3:
        d["teething"]      = pill_select("التسنين","Teething", TEETH_AR, "c_teeth")
    with c4:
        d["toilet"]        = pill_select("تدريب دورة المياه","Toilet training", TOILET_AR, "c_toilet")

    st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin:16px 0 8px;">النمو الحركي واللغوي / Motor & Speech</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        d["motor"]        = pill_select("النمو الحركي","Motor development", MOTOR_AR, "c_motor")
        d["motor_detail"] = ti("تفاصيل (مشي، جلوس...)","Motor details","c_motord")
    with c2:
        d["speech"]        = pill_select("الكلام","Speech", SPEECH_AR, "c_speech")
        d["speech_detail"] = ti("تفاصيل الكلام","Speech details","c_speechd")

    st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin:16px 0 8px;">القدرات المعرفية / Cognitive</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        d["attention"]     = pill_select("الانتباه","Attention",["طبيعي","ضعيف","لا ينطبق"],"c_attn")
    with c2:
        d["concentration"] = pill_select("التركيز","Concentration",["طبيعي","ضعيف","لا ينطبق"],"c_conc")
    with c3:
        d["comprehension"] = pill_select("الفهم والإدراك","Comprehension",["طبيعي","ضعيف","لا ينطبق"],"c_comp")
    d["dev_notes"] = ta("ملاحظات النمو","Developmental notes","c_devnotes",70)

    sec("👨‍👩‍👧","بيانات الأسرة","Family Details")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin-bottom:8px;">الأب / Father</div>', unsafe_allow_html=True)
        d["father_name"]       = ti("الاسم","Name","c_fn")
        d["father_age"]        = ti("السن","Age","c_fa")
        d["father_occ"]        = ti("الوظيفة","Occupation","c_fo")
        d["father_alive"]      = pill_select("الحالة","Status", ALIVE_M, "c_falive")
        d["father_hereditary"] = ti("مرض وراثي عند الأب","Hereditary disease","c_fh")
    with c2:
        st.markdown(f'<div style="font-weight:700;color:{DEEP_BLUE};font-family:Cairo;margin-bottom:8px;">الأم / Mother</div>', unsafe_allow_html=True)
        d["mother_name"]       = ti("الاسم","Name","c_mn")
        d["mother_age"]        = ti("السن","Age","c_ma")
        d["mother_occ"]        = ti("الوظيفة","Occupation","c_mo")
        d["mother_alive"]      = pill_select("الحالة","Status", ALIVE_F, "c_malive")
        d["mother_hereditary"] = ti("مرض وراثي عند الأم","Hereditary disease","c_mh")

    c1, c2 = st.columns(2)
    with c1:
        d["consanguinity"] = pill_select("القرابة بين الأبوين","Consanguinity", CONS_AR, "c_cons")
    with c2:
        d["parents_rel"]   = pill_select("العلاقة بين الأبوين","Parents relationship", PARENTS_REL, "c_prel")

    sec("👫","الإخوة والأخوات","Brothers & Sisters")
    siblings = []
    sib_cols = st.columns([1,2,1,2,2])
    for col, h in zip(sib_cols, ["النوع","الاسم","السن","التعليم","ملاحظات"]):
        col.markdown(f'<div class="sib-col-header">{h}</div>', unsafe_allow_html=True)

    for i in range(1, 6):
        c1,c2,c3,c4,c5 = st.columns([1,2,1,2,2])
        with c1: g  = st.selectbox("",SIB_GENDER,key=f"c_sg{i}",label_visibility="collapsed")
        with c2: n  = st.text_input("",key=f"c_sn{i}",placeholder=f"الاسم {i}",label_visibility="collapsed")
        with c3: ag = st.text_input("",key=f"c_sa{i}",placeholder="",label_visibility="collapsed")
        with c4: e  = st.selectbox("",["—"]+SIB_EDU,key=f"c_se{i}",label_visibility="collapsed")
        with c5: nt = st.text_input("",key=f"c_st{i}",placeholder="",label_visibility="collapsed")
        if n.strip():
            siblings.append({"gender":g,"name":n,"age":ag,"edu":e,"notes":nt})
    d["siblings"] = siblings

    c1, c2 = st.columns(2)
    with c1:
        d["sibling_rel"] = pill_select("علاقة الأخوة ببعض","Sibling relationship", SIB_REL, "c_sibrel")
    with c2:
        d["same_school"] = pill_select("هل الأخوة في نفس المدرسة؟","Same school?", SAME_SCH, "c_ssch")

    sec("🩺","الشكاوى وتاريخ المرض الحالي","Complaints & HPI")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["onset"]      = ti("متى بدأت الأعراض؟","Onset","c_onset")
    with c2:
        d["onset_mode"] = pill_select("طريقة البداية","Mode", ONSET_MODE, "c_omode")
    with c3:
        d["course"]     = pill_select("مسار المرض","Course", COURSE_AR, "c_course")
    d["complaints"] = ta("الشكاوى الرئيسية (C/O)","Chief Complaints","c_co",120)
    d["hpi"]        = ta("تاريخ المرض الحالي بالتفصيل (HPI)","History of Presenting Illness","c_hpi",220)

    sec("📋","التاريخ المرضي السابق","Past History")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["high_fever"]   = pill_select("حرارة ≥40 درجة؟","High fever?", YES_NO_NA, "c_hfever")
        d["head_trauma"]  = pill_select("ارتطام رأس؟","Head trauma?", YES_NO_NA, "c_htrauma")
        if st.session_state.get("c_htrauma") == "نعم":
            d["head_trauma_location"] = ti("مكان الارتطام","Location","c_htrauma_loc")
            d["head_trauma_details"]  = ti("كيف حدث؟","How?","c_htrauma_det")
        else:
            d["head_trauma_location"] = ""
            d["head_trauma_details"]  = ""
    with c2:
        d["convulsions"]  = pill_select("تشنجات؟","Convulsions?", YES_NO_NA, "c_conv")
        d["post_vaccine"] = pill_select("مضاعفات بعد التطعيم؟","Post-vaccine?", YES_NO_NA, "c_pvacc")
    with c3:
        d["prev_hosp"]    = pill_select("دخول مستشفى سابق؟","Previous hosp.?", YES_NO_NA, "c_phosp")
        d["prev_therapy"] = pill_select("جلسات علاجية سابقة؟","Previous therapy?", YES_NO_NA, "c_pther")
    d["past_hx"] = ta("تفاصيل التاريخ السابق","Details","c_past",100)

    sec("🧬","التاريخ العائلي","Family History")
    c1, c2 = st.columns(2)
    with c1:
        d["fam_psych"] = pill_select("مرض نفسي في الأسرة؟","Psychiatric?", YES_NO_NA, "c_fpsych")
        if st.session_state.get("c_fpsych") == "نعم":
            d["fam_psych_details"] = ti("تفاصيل","Details","c_fpsych_det")
        else:
            d["fam_psych_details"] = ""

        d["fam_neuro"] = pill_select("مرض عصبي في الأسرة؟","Neurological?", YES_NO_NA, "c_fneuro")
        if st.session_state.get("c_fneuro") == "نعم":
            d["fam_neuro_details"] = ti("تفاصيل","Details","c_fneuro_det")
        else:
            d["fam_neuro_details"] = ""
    with c2:
        d["fam_mr"] = pill_select("إعاقة ذهنية في الأسرة؟","MR in family?", YES_NO_NA, "c_fmr")
        if st.session_state.get("c_fmr") == "نعم":
            d["fam_mr_details"] = ti("تفاصيل","Details","c_fmr_det")
        else:
            d["fam_mr_details"] = ""

        d["fam_epilepsy"] = pill_select("صرع في الأسرة؟","Epilepsy?", YES_NO_NA, "c_fepil")
        if st.session_state.get("c_fepil") == "نعم":
            d["fam_epilepsy_details"] = ti("تفاصيل","Details","c_fepil_det")
        else:
            d["fam_epilepsy_details"] = ""
    d["family_hx"] = ta("تفاصيل التاريخ العائلي","Details","c_famhx",80)

    sec("🔬","الفحوصات","Investigations")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["had_ct"]   = pill_select("أشعة مقطعية CT؟","CT?", YES_NO_NA, "c_ct")
        d["had_mri"]  = pill_select("رنين مغناطيسي MRI؟","MRI?", YES_NO_NA, "c_mri")
    with c2:
        d["had_eeg"]  = pill_select("رسم مخ EEG؟","EEG?", YES_NO_NA, "c_eeg")
        d["had_iq"]   = pill_select("اختبار ذكاء SB5؟","IQ SB5?", YES_NO_NA, "c_iq")
    with c3:
        d["had_cars"] = pill_select("مقياس CARS؟","CARS?", YES_NO_NA, "c_cars")
        d["cars_score"]= ti("درجة CARS","CARS score","c_carsscore")
    d["investigations"] = ta("تفاصيل الفحوصات ونتائجها","Details","c_inv",80)

    sec("🏥","العمليات والجراحات","Operations & Surgeries")
    d["had_surg"]  = pill_select("عمليات جراحية سابقة؟","Previous surgeries?", YES_NO_NA, "c_hsurg")
    d["surgeries"] = ta("تفاصيل العمليات","Details","c_surg",60)

    sec("📊","التقييم السريري","Clinical Assessment")
    c1, c2 = st.columns(2)
    with c1:
        d["sleep"]           = pill_select("نمط النوم","Sleep", SLEEP_AR, "c_sleep")
        d["appetite"]        = pill_select("الشهية","Appetite", APPETITE_AR, "c_appetite")
        d["punishment"]      = pill_multi("طرق العقاب","Punishment", PUNISHMENT, "c_punish")
        d["stress_reaction"] = pill_multi("رد الفعل تجاه الضغوط","Stress reaction", STRESS_RXN, "c_stress")
    with c2:
        d["therapy"]         = ta("الجلسات العلاجية الحالية","Current therapy","c_therapy",100)
    d["extra_notes"] = ta("ملاحظات إضافية","Additional notes","c_extra",80)
    patient_name = d.get("name") or "الطفل"


# ══════════════════════════════════════════════════════════════
#  GENERATE BUTTON
# ══════════════════════════════════════════════════════════════
st.divider()
col_gen, col_sp = st.columns([2,3])
with col_gen:
    generate_clicked = st.button(
        "✦ توليد التقرير / Generate Report",
        type="primary",
        use_container_width=True
    )

if generate_clicked:
    # ── Build sibling text ──
    siblings = d.get("siblings", [])
    sib_text = "\n".join([
        f"  {i+1}. {sb['name']} | {sb['gender']} | السن: {sb['age']} | التعليم: {sb['edu']} | ملاحظات: {sb['notes'] or '—'}"
        for i, sb in enumerate(siblings)
    ]) or "لا يوجد إخوة مُدخَلون"

    # ── Build verbatim block (ACTUAL field values, not placeholder) ──
    verbatim_items = []
    long_text_fields = [
        ("الشكوى الرئيسية (C/O)", sv(d,'complaints')),
        ("تاريخ المرض الحالي (HPI)", sv(d,'hpi')),
        ("تفاصيل الحمل والولادة", sv(d,'preg_notes') if not is_adult else ""),
        ("تاريخ الأدوية - تفاصيل", sv(d,'drug_hx') if is_adult else ""),
        ("التاريخ المرضي السابق - تفاصيل", sv(d,'past_hx')),
        ("التاريخ العائلي - تفاصيل", sv(d,'family_hx')),
        ("الفحوصات - تفاصيل", sv(d,'investigations')),
        ("ملاحظات النمو", sv(d,'dev_notes') if not is_adult else ""),
        ("الجلسات العلاجية الحالية", sv(d,'therapy') if not is_adult else ""),
        ("ملاحظات إضافية", sv(d,'extra_notes')),
    ]
    for heading, text in long_text_fields:
        if text and text != "لم يُذكر":
            verbatim_items.append((heading, text))

    # ── Data block for AI ──
    if is_adult:
        data_block = f"""
المريض: {sv(d,'name')} | تاريخ الميلاد: {sv(d,'birthdate')} | السن: {sv(d,'age')} | النوع: {sv(d,'gender')}
التاريخ: {sv(d,'date')} | الأخصائي: {history_by or 'لم يُذكر'} | نوع التاريخ: {sv(d,'htype')}
الهاتف: {sv(d,'phone')} | مصدر الإحالة: {sv(d,'referral')}
الوظيفة: {sv(d,'occupation')} — {sv(d,'occ_detail')} | التعليم: {sv(d,'education')}
الحالة الاجتماعية: {sv(d,'social')} | التدخين: {sv(d,'smoking')} | الهوايات: {sv(d,'hobbies')}

الأب: {sv(d,'father_name')} | السن: {sv(d,'father_age')} | الوظيفة: {sv(d,'father_occ')} | الحالة: {sv(d,'father_alive')}
الأم: {sv(d,'mother_name')} | السن: {sv(d,'mother_age')} | الوظيفة: {sv(d,'mother_occ')} | الحالة: {sv(d,'mother_alive')}
القرابة: {sv(d,'consanguinity')} | يعيشان معاً: {sv(d,'parents_together')} | مرض مزمن: {sv(d,'chronic')}

الزوج/الزوجة: {sv(d,'spouse_name')} | السن: {sv(d,'spouse_age')} | الوظيفة: {sv(d,'spouse_occ')}
مدة الزواج: {sv(d,'marriage_dur')} | الخطوبة: {sv(d,'engagement')} | كتب كتاب: {sv(d,'katb')}
جودة الزواج: {sv(d,'marriage_qual')} | العلاقة قبل الزواج: {sv(d,'pre_marriage')} | الأبناء: {sv(d,'num_children')}

الإخوة:\n{sib_text}

الأعراض: بداية: {sv(d,'onset')} | طريقة: {sv(d,'onset_mode')} | مسار: {sv(d,'course')}
الشكاوى: {sv(d,'complaints')}
HPI: {sv(d,'hpi')}

الأدوية: {sv(d,'on_meds')} | الالتزام: {sv(d,'compliance')} | تفاصيل: {sv(d,'drug_hx')}
التاريخ السابق: نفسي: {sv(d,'prev_psych')} | مستشفى: {sv(d,'prev_hosp')} | تفاصيل: {sv(d,'past_hx')}
التاريخ العائلي: نفسي: {sv(d,'fam_psych')}{(' — '+sv(d,'fam_psych_details')) if d.get('fam_psych_details') else ''} | عصبي: {sv(d,'fam_neuro')}{(' — '+sv(d,'fam_neuro_details')) if d.get('fam_neuro_details') else ''} | تفاصيل: {sv(d,'family_hx')}
الفحوصات: {sv(d,'had_inv')} | {sv(d,'investigations')}
الجراحات: {sv(d,'had_surg')} | {sv(d,'surgeries')}
النوم: {sv(d,'sleep')} | الشهية: {sv(d,'appetite')} | انتحارية: {sv(d,'suicidal')} | بصيرة: {sv(d,'insight')}
مواد: {sv(d,'substance')} — {sv(d,'substance_details')}
ملاحظات: {sv(d,'extra_notes')}
"""
    else:
        data_block = f"""
الطفل: {sv(d,'name')} | تاريخ الميلاد: {sv(d,'birthdate')} | السن: {sv(d,'age')} | النوع: {sv(d,'gender')}
التاريخ: {sv(d,'date')} | الأخصائي: {history_by or 'لم يُذكر'} | الهاتف: {sv(d,'phone')}
يعيش مع: {sv(d,'lives_with')} | المدرسة: {sv(d,'school')} | الصف: {sv(d,'grade')} | المستوى: {sv(d,'academic')}
ترتيب الميلاد: {sv(d,'birth_order')} | وقت الشاشة: {sv(d,'screen_time')}
مرغوب فيه: {sv(d,'wanted')} | النوع المرغوب: {sv(d,'gender_des')}

الحمل: {sv(d,'pregnancy')} | نوع الولادة: {sv(d,'birth_type')} | مضاعفات: {sv(d,'birth_comp')}
التطعيمات: {sv(d,'vacc_status')} | مضاعفات التطعيم: {sv(d,'vacc_comp')}
ملاحظات الحمل: {sv(d,'preg_notes')}
الرضاعة: {sv(d,'breastfeeding')} | الفطام: {sv(d,'weaning')} | التسنين: {sv(d,'teething')} | الحمام: {sv(d,'toilet')}
الحركة: {sv(d,'motor')} — {sv(d,'motor_detail')} | الكلام: {sv(d,'speech')} — {sv(d,'speech_detail')}
الانتباه: {sv(d,'attention')} | التركيز: {sv(d,'concentration')} | الفهم: {sv(d,'comprehension')}
ملاحظات النمو: {sv(d,'dev_notes')}

الأب: {sv(d,'father_name')} | السن: {sv(d,'father_age')} | الوظيفة: {sv(d,'father_occ')} | الحالة: {sv(d,'father_alive')} | وراثي: {sv(d,'father_hereditary')}
الأم: {sv(d,'mother_name')} | السن: {sv(d,'mother_age')} | الوظيفة: {sv(d,'mother_occ')} | الحالة: {sv(d,'mother_alive')} | وراثي: {sv(d,'mother_hereditary')}
القرابة: {sv(d,'consanguinity')} | علاقة الأبوين: {sv(d,'parents_rel')}

الإخوة:\n{sib_text}
علاقة الأخوة: {sv(d,'sibling_rel')} | نفس المدرسة: {sv(d,'same_school')}

الأعراض: بداية: {sv(d,'onset')} | طريقة: {sv(d,'onset_mode')} | مسار: {sv(d,'course')}
الشكاوى: {sv(d,'complaints')}
HPI: {sv(d,'hpi')}

التاريخ السابق: حرارة: {sv(d,'high_fever')} | رأس: {sv(d,'head_trauma')}{(' — '+sv(d,'head_trauma_location')+' — '+sv(d,'head_trauma_details')) if sv(d,'head_trauma_location','') != 'لم يُذكر' else ''}
تشنجات: {sv(d,'convulsions')} | تطعيم: {sv(d,'post_vaccine')} | مستشفى: {sv(d,'prev_hosp')} | جلسات: {sv(d,'prev_therapy')}
تفاصيل: {sv(d,'past_hx')}
التاريخ العائلي: نفسي: {sv(d,'fam_psych')}{(' — '+sv(d,'fam_psych_details')) if d.get('fam_psych_details') else ''} | عصبي: {sv(d,'fam_neuro')}{(' — '+sv(d,'fam_neuro_details')) if d.get('fam_neuro_details') else ''} | MR: {sv(d,'fam_mr')}{(' — '+sv(d,'fam_mr_details')) if d.get('fam_mr_details') else ''} | صرع: {sv(d,'fam_epilepsy')}{(' — '+sv(d,'fam_epilepsy_details')) if d.get('fam_epilepsy_details') else ''}
تفاصيل: {sv(d,'family_hx')}
CT: {sv(d,'had_ct')} | MRI: {sv(d,'had_mri')} | EEG: {sv(d,'had_eeg')} | SB5: {sv(d,'had_iq')} | CARS: {sv(d,'had_cars')} درجة: {sv(d,'cars_score')}
تفاصيل الفحوصات: {sv(d,'investigations')}
الجراحات: {sv(d,'had_surg')} — {sv(d,'surgeries')}
النوم: {sv(d,'sleep')} | الشهية: {sv(d,'appetite')} | عقاب: {sv(d,'punishment')} | رد فعل: {sv(d,'stress_reaction')}
الجلسات الحالية: {sv(d,'therapy')}
ملاحظات: {sv(d,'extra_notes')}
"""

    # ── Build verbatim section text ──
    verbatim_section_en = ""
    verbatim_section_ar = ""
    if verbatim_items:
        verbatim_section_en = "\n\nORIGINAL ARABIC RESPONSES (verbatim — do not translate)\n"
        verbatim_section_en += "=" * 50 + "\n"
        for heading, text in verbatim_items:
            verbatim_section_en += f"\n{heading}:\n{text}\n"

        verbatim_section_ar = "\n\nالنصوص الأصلية كما أُدخلت\n"
        verbatim_section_ar += "=" * 50 + "\n"
        for heading, text in verbatim_items:
            verbatim_section_ar += f"\n{heading}:\n{text}\n"
    else:
        verbatim_section_en = "\n\nORIGINAL ARABIC RESPONSES\nNo long text responses were provided.\n"
        verbatim_section_ar = "\n\nالنصوص الأصلية\nلم يتم إدخال نصوص طويلة.\n"

    # ── ENGLISH PROMPT ──
    en_prompt = f"""You are a clinical report formatter. Generate a compact professional clinical history report in ENGLISH only.

STRICT RULES:
1. Write sections 1–8 in English only. No Arabic words anywhere in sections 1–8.
2. Convert all Arabic MCQ answers to English equivalents (e.g. "ذكر" → Male, "مفاجئ" → Sudden).
3. Omit any field that is "لم يُذكر" (not reported). Skip entire sections if empty.
4. "No" / "لا" answers: omit unless clinically significant.
5. No diagnosis, interpretation, or assumptions beyond what is stated.
6. No markdown symbols (**, ##, ---, bullets with -).
7. Section titles: ALL CAPS numbered. Example: 1. PATIENT INFORMATION
8. Table rows: Field | Value  (pipe format)
9. Keep compact: group short fields inline where logical.
10. The final section must contain the original Arabic long-text responses EXACTLY as written — no translation, no modification.

STRUCTURE:

CLINICAL HISTORY REPORT
Patient | {sv(d,'name')}
Form Type | {"Adult" if is_adult else "Child"}
Specialist | {history_by or 'Not recorded'}
Date | {sv(d,'date')}  |  Phone | {sv(d,'phone')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

CLINICAL SUMMARY
Write 2–4 sentences summarizing: patient identity, chief complaint, key background. No diagnosis.

1. PATIENT INFORMATION
Compact table of all demographic and personal fields.

2. PRESENTING CONCERNS
Onset | [value]
Mode | [value]
Course | [value]
Then list chief symptoms as plain lines.

3. {"FAMILY & MARRIAGE" if is_adult else "FAMILY BACKGROUND"}
Compact tables for parents, {"marriage," if is_adult else ""} siblings, consanguinity.

{"4. DEVELOPMENTAL HISTORY" if not is_adult else ""}
{"Compact tables: Prenatal/natal, feeding/growth, motor/speech, cognitive." if not is_adult else ""}

{"4. PAST HISTORY" if is_adult else "5. PAST HISTORY"}
Compact table of past medical history.

{"5. FAMILY HISTORY" if is_adult else "6. FAMILY HISTORY"}
Compact table.

{"6. MEDICATIONS & COMPLIANCE" if is_adult else ""}
{"Compact table." if is_adult else ""}

{"6. CLINICAL OBSERVATIONS" if not is_adult else "7. CLINICAL OBSERVATIONS"}
Inline format: Sleep | Appetite | {"Suicidal ideation | Insight | Substance use" if is_adult else "Attention | Concentration | Comprehension | Punishment | Stress reaction"}

{"7. INVESTIGATIONS" if is_adult else "8. INVESTIGATIONS"}
Compact table. Skip if none.

{"8. ORIGINAL ARABIC RESPONSES" if is_adult else "9. ORIGINAL ARABIC RESPONSES"}
{verbatim_section_en}

DATA:
{data_block}
"""

    # ── ARABIC PROMPT ──
    ar_prompt = f"""أنت مُنسِّق تقارير سريرية. أنشئ تقريراً سريرياً موجزاً ومهنياً باللغة العربية فقط.

قواعد صارمة:
1. اكتب التقرير كاملاً بالعربية. لا تستخدم الإنجليزية إلا في الاختصارات الطبية (CT, MRI, EEG, IQ, CARS, HPI, C/O).
2. احذف أي حقل قيمته "لم يُذكر". احذف الأقسام الكاملة إذا كانت فارغة.
3. إجابات "لا": احذفها إلا إذا كانت ذات أهمية سريرية.
4. لا تضف أي تشخيص أو تفسير أو معلومات غير موجودة في البيانات.
5. لا تستخدم رموز markdown (**, ##, ---).
6. عناوين الأقسام: أرقام + حروف كبيرة. مثال: ١. البيانات الشخصية
7. الجداول: الحقل | القيمة
8. مختصر وموجز: اجمع الحقول القصيرة في سطر واحد.
9. اتجاه النص: من اليمين إلى اليسار.
10. القسم الأخير: انسخ النصوص الأصلية الطويلة كما أُدخلت بالضبط دون تعديل.

الهيكل:

تقرير التاريخ المرضي السريري
المريض | {sv(d,'name')}
نوع الاستمارة | {"بالغ" if is_adult else "طفل"}
الأخصائي | {history_by or 'غير مسجل'}
التاريخ | {sv(d,'date')}  |  الهاتف | {sv(d,'phone')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

ملخص سريري
اكتب ٢–٤ جمل تلخيصية: هوية المريض، الشكوى الرئيسية، الخلفية الأساسية. بدون تشخيص.

١. البيانات الشخصية
جدول موجز لجميع الحقول الديموغرافية.

٢. الشكاوى وتاريخ المرض
البداية | [القيمة]
طريقة البداية | [القيمة]
المسار | [القيمة]
ثم قائمة بالأعراض الرئيسية.

٣. {"بيانات الأسرة والزواج" if is_adult else "بيانات الأسرة"}
جداول موجزة للأبوين {"والزواج" if is_adult else ""} والإخوة والقرابة.

{"٤. مراحل النمو" if not is_adult else ""}
{"جداول موجزة: الحمل والولادة، التغذية والنمو، الحركة والكلام، المعرفي." if not is_adult else ""}

{"٤. التاريخ المرضي السابق" if is_adult else "٥. التاريخ المرضي السابق"}
جدول موجز.

{"٥. التاريخ العائلي" if is_adult else "٦. التاريخ العائلي"}
جدول موجز.

{"٦. الأدوية والالتزام" if is_adult else ""}
{"جدول موجز." if is_adult else ""}

{"٦. التقييم السريري" if not is_adult else "٧. التقييم السريري"}
نمط النوم | الشهية | {"الأفكار الانتحارية | البصيرة | المواد" if is_adult else "الانتباه | التركيز | الفهم | العقاب | رد الفعل"}

{"٧. الفحوصات" if is_adult else "٨. الفحوصات"}
جدول موجز إن وجد.

{"٨. النصوص الأصلية" if is_adult else "٩. النصوص الأصلية"}
{verbatim_section_ar}

البيانات:
{data_block}
"""

    with st.spinner("⏳ جاري إنشاء التقارير..."):
        try:
            client  = Groq(api_key=st.secrets["GROQ_API_KEY"])

            resp_en = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role":"user","content":en_prompt}],
                max_tokens=4000
            )
            resp_ar = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role":"user","content":ar_prompt}],
                max_tokens=4000
            )

            st.session_state["report_en"]    = resp_en.choices[0].message.content
            st.session_state["report_ar"]    = resp_ar.choices[0].message.content
            st.session_state["report_pname"] = patient_name
            st.session_state["report_sheet"] = "بالغ" if is_adult else "طفل"
            st.session_state["report_by"]    = history_by or "—"
            st.session_state["verbatim_items"] = verbatim_items

        except Exception as e:
            st.error(f"خطأ في توليد التقرير: {str(e)}")


# ══════════════════════════════════════════════════════════════
#  DISPLAY + DOWNLOAD
# ══════════════════════════════════════════════════════════════
if st.session_state.get("report_en"):
    rt_en   = st.session_state["report_en"]
    rt_ar   = st.session_state["report_ar"]
    pn      = st.session_state.get("report_pname","المريض")
    rs      = st.session_state.get("report_sheet","")
    rb_     = st.session_state.get("report_by","—")
    vb      = st.session_state.get("verbatim_items",[])
    fn_en   = f"{pn.replace(' ','_')}_EN_History.docx"
    fn_ar   = f"{pn.replace(' ','_')}_AR_History.docx"

    st.divider()
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,{DEEP_BLUE},{MID_BLUE});
                border-radius:12px;padding:16px 24px;margin-bottom:20px;color:white;
                font-family:'Cairo',sans-serif;font-size:15px;font-weight:600;">
        ✅ تم إنشاء التقريرين بنجاح — {pn}
    </div>
    """, unsafe_allow_html=True)

    tab_en, tab_ar = st.tabs(["🇬🇧 English Report", "🇸🇦 التقرير العربي"])
    with tab_en:
        st.text_area("", value=rt_en, height=500, label_visibility="collapsed")
    with tab_ar:
        st.text_area("", value=rt_ar, height=500, label_visibility="collapsed")

    # ── DOCX Builder ──────────────────────────────────────────
    def build_docx(report_text: str, lang: str) -> io.BytesIO:
        """Build a premium Word document for either 'en' or 'ar' report."""
        is_rtl = (lang == "ar")
        doc = Document()

        # Page setup
        for sec_ in doc.sections:
            sec_.top_margin    = Cm(2.0)
            sec_.bottom_margin = Cm(2.0)
            sec_.left_margin   = Cm(2.2)
            sec_.right_margin  = Cm(2.2)

        # Page border
        for sec_ in doc.sections:
            sp = sec_._sectPr
            pb = OxmlElement('w:pgBorders')
            pb.set(qn('w:offsetFrom'),'page')
            for side in ('top','left','bottom','right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),'single')
                b.set(qn('w:sz'),'12')
                b.set(qn('w:space'),'24')
                b.set(qn('w:color'),'3B5FC0')
                pb.append(b)
            sp.append(pb)

        # Footer with page numbers
        for sec_ in doc.sections:
            ft = sec_.footer
            fp = ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
            fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = fp.add_run()
            r.font.size = Pt(9)
            r.font.color.rgb = CLINIC_BLUE_RGB
            for tag, text in [('begin',None),(None,' PAGE '),('end',None)]:
                if tag:
                    el = OxmlElement('w:fldChar')
                    el.set(qn('w:fldCharType'), tag)
                    r._r.append(el)
                else:
                    it = OxmlElement('w:instrText')
                    it.text = text
                    r._r.append(it)

        def set_rtl(paragraph):
            if is_rtl:
                pPr = paragraph._p.get_or_add_pPr()
                b = OxmlElement("w:bidi")
                pPr.append(b)
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "right")
                pPr.append(jc)

        def add_para(text, bold=False, size=11, color=None,
                     space_before=0, space_after=4, italic=False,
                     alignment=None, keep_next=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            if keep_next:
                p.paragraph_format.keep_with_next = True
            set_rtl(p)
            if alignment:
                p.alignment = alignment
            r = p.add_run(text)
            r.font.size   = Pt(size)
            r.font.name   = "Arial"
            r.font.bold   = bold
            r.font.italic = italic
            if color:
                r.font.color.rgb = color
            return p

        def add_section_title(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.keep_with_next = True
            set_rtl(p)
            r = p.add_run(text.strip())
            r.font.size  = Pt(13)
            r.font.name  = "Arial"
            r.font.bold  = True
            r.font.color.rgb = CLINIC_BLUE_RGB
            # Underline via bottom border
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),'single')
            bot.set(qn('w:sz'),'6')
            bot.set(qn('w:space'),'2')
            bot.set(qn('w:color'),'3B5FC0')
            pBdr.append(bot)
            pPr.append(pBdr)

        def make_table():
            t = doc.add_table(rows=0, cols=2)
            t.style = 'Table Grid'
            try:
                tPr = t._tbl.tblPr
                tW  = OxmlElement('w:tblW')
                tW.set(qn('w:w'),'9026')
                tW.set(qn('w:type'),'dxa')
                tPr.append(tW)
                tg = OxmlElement('w:tblGrid')
                for w in [3000, 6026]:
                    gc = OxmlElement('w:gridCol')
                    gc.set(qn('w:w'), str(w))
                    tg.append(gc)
                t._tbl.insert(0, tg)
            except: pass
            return t

        def add_row(table, field, value, header=False):
            row  = table.add_row()
            trPr = row._tr.get_or_add_trPr()
            cs   = OxmlElement('w:cantSplit')
            cs.set(qn('w:val'),'1')
            trPr.append(cs)

            for idx, (cell, txt, bold_) in enumerate([
                (row.cells[0], field, True),
                (row.cells[1], value, header)
            ]):
                cell.text = ""
                p = cell.paragraphs[0]
                if is_rtl:
                    pPr = p._p.get_or_add_pPr()
                    pPr.append(OxmlElement("w:bidi"))
                    jc = OxmlElement("w:jc")
                    jc.set(qn("w:val"),"right")
                    pPr.append(jc)

                # Multi-line value support
                lines = str(txt).split('\n') if '\n' in str(txt) else [str(txt)]
                for li, line in enumerate(lines):
                    vp = p if li == 0 else cell.add_paragraph()
                    vr = vp.add_run(line.strip())
                    vr.font.size = Pt(10)
                    vr.font.name = "Arial"
                    vr.font.bold = bold_
                    if header:
                        vr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

                tc  = cell._tc
                tcP = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),'clear')
                shd.set(qn('w:color'),'auto')
                if header:
                    shd.set(qn('w:fill'), '3B5FC0' if idx==0 else '5B7FD0')
                elif idx == 0:
                    shd.set(qn('w:fill'), 'E8EEFF')
                else:
                    shd.set(qn('w:fill'), 'FFFFFF')
                tcP.append(shd)
                mg = OxmlElement('w:tcMar')
                for side in ['top','bottom','left','right']:
                    m = OxmlElement(f'w:{side}')
                    m.set(qn('w:w'),'60')
                    m.set(qn('w:type'),'dxa')
                    mg.append(m)
                tcP.append(mg)

        # ── Document header ──
        # Logo + title row
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(0)
        p_hdr.paragraph_format.space_after  = Pt(6)
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if os.path.exists(LOGO_PATH):
            p_hdr.add_run().add_picture(LOGO_PATH, width=Inches(1.1))

        title_text = "Clinical History Report" if lang=="en" else "تقرير التاريخ المرضي السريري"
        r_title = p_hdr.add_run(f"   {title_text}")
        r_title.font.name  = "Arial"
        r_title.font.size  = Pt(18)
        r_title.font.bold  = True
        r_title.font.color.rgb = CLINIC_BLUE_RGB

        # Subtitle
        sub_text = DOCTOR["title1"] if lang=="en" else "عيادة د. هاني الحناوي — طب وجراحة الأعصاب والنفس"
        add_para(sub_text, size=10, color=LOGO_BLUE_RGB,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

        # Separator line
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(2)
        p_sep.paragraph_format.space_after  = Pt(8)
        pPr = p_sep._p.get_or_add_pPr()
        pBdr2 = OxmlElement('w:pBdr')
        bot2  = OxmlElement('w:bottom')
        bot2.set(qn('w:val'),'single'); bot2.set(qn('w:sz'),'8')
        bot2.set(qn('w:space'),'2');    bot2.set(qn('w:color'),'A1B6F3')
        pBdr2.append(bot2); pPr.append(pBdr2)

        # ── Parse report lines ──
        in_table      = False
        current_table = None
        lines         = report_text.split('\n')

        # Section patterns
        sec_en = re.compile(r'^\d+\.\s+[A-Z][A-Z\s&/\(\)]+$')
        sec_ar = re.compile(r'^[١٢٣٤٥٦٧٨٩\d]+[\.،]\s+[\u0600-\u06FF]')
        header_words = {'CLINICAL HISTORY REPORT','CLINICAL SUMMARY','REPORT HEADER',
                       'تقرير التاريخ المرضي','ملخص سريري'}

        for line in lines:
            ls = line.strip()
            if not ls:
                if in_table:
                    in_table = False
                    current_table = None
                    doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue

            # Section titles
            is_section = (sec_en.match(ls) or sec_ar.match(ls) or
                         ls in header_words or ls.upper() in header_words)
            if is_section:
                in_table = False; current_table = None
                add_section_title(ls)
                continue

            # Separator lines
            if ls.startswith('━') or ls.startswith('═') or ls.startswith('---'):
                in_table = False; current_table = None
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(4)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                b    = OxmlElement('w:bottom')
                b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
                b.set(qn('w:space'),'1');    b.set(qn('w:color'),'C5D3F5')
                pBdr.append(b); pPr.append(pBdr)
                continue

            # Table rows with pipe
            if '|' in ls:
                parts = [p.strip() for p in ls.split('|') if p.strip()]
                if not parts:
                    continue
                # Skip markdown separator rows (e.g. |---|---|)
                if all(set(p) <= set('-: ') for p in parts):
                    continue
                # Skip generic header rows
                skip_pairs = [("field","value"),("الحقل","القيمة"),("milestone","finding")]
                if (len(parts) >= 2 and
                    (parts[0].strip('* ').lower(), parts[1].strip('* ').lower()) in skip_pairs):
                    continue

                if not in_table or current_table is None:
                    in_table = True
                    current_table = make_table()
                    add_row(current_table, "Field" if lang=="en" else "الحقل",
                            "Details" if lang=="en" else "التفاصيل", header=True)

                if len(parts) >= 2:
                    field = parts[0].strip('* ')
                    value = ' | '.join(parts[1:])
                    add_row(current_table, field, value)
                elif len(parts) == 1:
                    add_row(current_table, parts[0].strip('* '), '')
                continue

            # Arabic verbatim section headings
            is_ar_text = any('\u0600' <= c <= '\u06ff' for c in ls)
            if ls.endswith(':') and is_ar_text and len(ls) < 60:
                in_table = False; current_table = None
                add_para(ls.rstrip(':'), bold=True, size=11,
                         color=DARK_BLUE_RGB, space_before=10, space_after=2,
                         keep_next=True)
                continue

            # Regular paragraph
            in_table = False; current_table = None
            add_para(ls, size=10.5, space_before=0, space_after=3)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ── Render download buttons ──
    st.markdown("---")
    dl_col1, dl_col2, dl_col3 = st.columns(3)

    with dl_col1:
        en_buf = build_docx(rt_en, "en")
        st.download_button(
            "📄 Download English Report",
            data=en_buf,
            file_name=fn_en,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with dl_col2:
        ar_buf = build_docx(rt_ar, "ar")
        st.download_button(
            "📄 تحميل التقرير العربي",
            data=ar_buf,
            file_name=fn_ar,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with dl_col3:
        if st.button("📧 إرسال بالبريد / Send Email", use_container_width=True):
            try:
                en_buf2 = build_docx(rt_en, "en")
                ar_buf2 = build_docx(rt_ar, "ar")
                msg = MIMEMultipart()
                msg['From']    = GMAIL_USER
                msg['To']      = RECIPIENT_EMAIL
                msg['Subject'] = f"تقرير التاريخ المرضي — {pn} ({rs})"
                body = f"المريض: {pn}\nالنوع: {rs}\nالأخصائي: {rb_}\n\nمرفق التقريران (عربي وإنجليزي)."
                msg.attach(MIMEText(body, 'plain', 'utf-8'))
                for buf_, fname_ in [(en_buf2, fn_en), (ar_buf2, fn_ar)]:
                    part = MIMEBase('application','octet-stream')
                    part.set_payload(buf_.read())
                    encoders.encode_base64(part)
                    part.add_header('Content-Disposition', f'attachment; filename="{fname_}"')
                    msg.attach(part)
                with smtplib.SMTP_SSL('smtp.gmail.com', 465) as srv:
                    srv.login(GMAIL_USER, GMAIL_PASS)
                    srv.sendmail(GMAIL_USER, RECIPIENT_EMAIL, msg.as_string())
                st.success(f"✅ تم الإرسال إلى {RECIPIENT_EMAIL}")
            except Exception as e:
                st.error(f"خطأ في الإرسال: {str(e)}")
