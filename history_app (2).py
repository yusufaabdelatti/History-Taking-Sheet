import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, HRFlowable, Image as RLImage
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import io, os, re, smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date

# ══════════════════════════════════════════════════════════════
#  CONFIGURATION
# ══════════════════════════════════════════════════════════════

RECIPIENT_EMAIL = "yusuf.a.abdelatti@gmail.com"
GMAIL_USER      = "yusuf.a.abdelatti@gmail.com"
GMAIL_PASS      = "erjl ehlj wpyg mfgx"
LOGO_FILE       = "logo.png"

CLINIC_BLUE = RGBColor(0x1A, 0x5C, 0xB8)
DOCTOR = {
    "name":    "Dr. Hany Elhennawy",
    "title1":  "Consultant of Neuro-Psychiatry",
    "title2":  "Aviation Medical Council — Faculty of Medicine, 6th October University",
    "title3":  "MD of Neuroscience Research, Karolinska Institute — Sweden",
    "title4":  "Member of I.S.N.R",
    "address": "16 Hesham Labib St., off Makram Ebeid St. Ext., next to Mobilia Saad Mohamed Saad",
    "phone":   "+20 1000756200",
}

# ══════════════════════════════════════════════════════════════
#  PAGE CONFIG & CSS
# ══════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="التاريخ المرضي — د. هاني الحناوي",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@300;400;500;600&family=Jost:wght@300;400;500&display=swap');

:root {
    --cream: #F7F3EE;
    --white: #FFFFFF;
    --deep: #1C1917;
    --warm: #8B7355;
    --accent: #C4956A;
    --blue: #1A5CB8;
    --blue-light: #E8F0FE;
    --border: #DDD5C8;
    --selected: #2D2926;
    --red: #B71C1C;
    --green: #2E7D32;
}

html, body, [class*="css"] {
    font-family: 'Jost', sans-serif;
    background-color: var(--cream);
    color: var(--deep);
}
.stApp { background-color: var(--cream); }

/* ── Page header ── */
.page-header {
    text-align: center;
    padding: 2rem 0 1.5rem 0;
    border-bottom: 1px solid var(--border);
    margin-bottom: 2rem;
}
.page-header h1 {
    font-family: 'Cormorant Garamond', serif;
    font-size: 2.2rem;
    font-weight: 400;
    color: var(--blue);
    margin-bottom: 0.2rem;
    letter-spacing: 0.02em;
}
.page-header p {
    color: var(--warm);
    font-size: 0.82rem;
    letter-spacing: 0.1em;
    text-transform: uppercase;
}

/* ── Section cards ── */
.section-card {
    background: var(--white);
    border: 1px solid var(--border);
    border-top: 3px solid var(--blue);
    border-radius: 4px;
    padding: 1.6rem 1.8rem 1rem 1.8rem;
    margin-bottom: 1.4rem;
}
.section-title {
    font-family: 'Cormorant Garamond', serif;
    font-size: 1.3rem;
    font-weight: 500;
    color: var(--blue);
    margin-bottom: 1rem;
    padding-bottom: 0.5rem;
    border-bottom: 1px solid var(--blue-light);
    letter-spacing: 0.02em;
}
.section-title-ar {
    font-size: 0.8rem;
    color: var(--warm);
    font-family: 'Jost', sans-serif;
    font-weight: 300;
    margin-left: 0.5rem;
}

/* ── Field labels ── */
.field-label {
    font-size: 0.82rem;
    font-weight: 500;
    color: var(--deep);
    margin-bottom: 0.3rem;
    letter-spacing: 0.01em;
}
.field-label-ar {
    font-size: 0.78rem;
    color: var(--warm);
    font-weight: 300;
}

/* ── Radio buttons → dots ── */
div[data-testid="stRadio"] > label { display: none !important; }
div[data-testid="stRadio"] > div {
    gap: 0.4rem !important;
    flex-wrap: wrap !important;
}
div[data-testid="stRadio"] > div > label {
    background: var(--cream) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 20px !important;
    padding: 0.38rem 0.9rem !important;
    cursor: pointer !important;
    font-size: 0.82rem !important;
    color: var(--deep) !important;
    font-family: 'Jost', sans-serif !important;
    font-weight: 400 !important;
    transition: all 0.15s ease !important;
    white-space: nowrap !important;
}
div[data-testid="stRadio"] > div > label:hover {
    border-color: var(--blue) !important;
    background: var(--blue-light) !important;
    color: var(--blue) !important;
}
div[data-testid="stRadio"] > div > label[data-baseweb="radio"] {
    border-color: var(--blue) !important;
    background: var(--blue) !important;
    color: white !important;
}

/* ── Selectbox ── */
div[data-testid="stSelectbox"] > div > div {
    background: var(--white) !important;
    border: 1px solid var(--border) !important;
    border-radius: 3px !important;
    font-family: 'Jost', sans-serif !important;
    font-size: 0.85rem !important;
}

/* ── Multiselect ── */
div[data-testid="stMultiSelect"] > div {
    background: var(--white) !important;
    border: 1px solid var(--border) !important;
    border-radius: 3px !important;
}
div[data-testid="stMultiSelect"] span[data-baseweb="tag"] {
    background: var(--blue-light) !important;
    color: var(--blue) !important;
    border-radius: 12px !important;
    font-size: 0.78rem !important;
}

/* ── Text inputs ── */
div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea {
    background: var(--white) !important;
    border: 1px solid var(--border) !important;
    border-radius: 3px !important;
    font-family: 'Jost', sans-serif !important;
    font-size: 0.88rem !important;
    color: var(--deep) !important;
}
div[data-testid="stTextInput"] input:focus,
div[data-testid="stTextArea"] textarea:focus {
    border-color: var(--blue) !important;
    box-shadow: 0 0 0 2px rgba(26,92,184,0.1) !important;
}

/* ── Buttons ── */
.stButton > button {
    font-family: 'Jost', sans-serif !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    border-radius: 2px !important;
    padding: 0.65rem 1.8rem !important;
    transition: all 0.2s ease !important;
}
.stButton > button[kind="primary"] {
    background: var(--blue) !important;
    color: white !important;
    border: none !important;
}
.stButton > button[kind="primary"]:hover {
    background: #1445A0 !important;
}
.stButton > button[kind="secondary"] {
    background: var(--white) !important;
    color: var(--blue) !important;
    border: 1px solid var(--blue) !important;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: var(--white) !important;
    border-right: 1px solid var(--border) !important;
}
[data-testid="stSidebar"] h2 {
    font-family: 'Cormorant Garamond', serif !important;
    color: var(--blue) !important;
    font-size: 1.1rem !important;
}

/* ── Caption ── */
.stCaption { color: var(--warm) !important; font-size: 0.78rem !important; }

/* ── Divider ── */
hr { border-color: var(--border) !important; }

/* ── Age badge ── */
.age-badge {
    display: inline-block;
    background: var(--blue-light);
    color: var(--blue);
    border-radius: 12px;
    padding: 2px 10px;
    font-size: 0.78rem;
    font-weight: 500;
    margin-top: 4px;
}

/* ── Generate button ── */
.generate-btn .stButton > button {
    background: var(--blue) !important;
    color: white !important;
    border: none !important;
    width: 100% !important;
    padding: 1rem 2rem !important;
    font-size: 0.9rem !important;
}

/* ── Report output ── */
.report-box {
    background: var(--white);
    border: 1px solid var(--border);
    border-top: 3px solid var(--blue);
    border-radius: 4px;
    padding: 1.5rem;
    margin-top: 1rem;
}

/* ── Siblings row ── */
.sibling-row {
    background: var(--cream);
    border: 1px solid var(--border);
    border-radius: 4px;
    padding: 0.8rem 1rem;
    margin-bottom: 0.6rem;
}

/* ── Info box ── */
.info-box {
    background: var(--blue-light);
    border-left: 3px solid var(--blue);
    padding: 0.8rem 1rem;
    border-radius: 0 4px 4px 0;
    font-size: 0.85rem;
    color: var(--blue);
    margin: 0.5rem 0 1rem 0;
}

.warning-box {
    background: #FFF8F0;
    border-left: 3px solid #E07B39;
    padding: 0.8rem 1rem;
    border-radius: 0 4px 4px 0;
    font-size: 0.85rem;
    color: #7A3D1A;
    margin: 0.5rem 0;
}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  CHOICE LISTS
# ══════════════════════════════════════════════════════════════

NA = "— Select —"
YES_NO_NA   = ["Yes / نعم", "No / لا", "N/A / لا ينطبق"]
YES_NO      = ["Yes / نعم", "No / لا"]
GENDER      = ["Male / ذكر", "Female / أنثى"]
EDU         = [NA,"Illiterate","Primary","Preparatory","Secondary","University","Postgraduate"]
OCC         = [NA,"Government Employee","Private Sector","Self-employed","Student","Housewife","Retired","Unemployed","Other"]
SOCIAL      = [NA,"Single","Married","Divorced","Widowed","Separated"]
SMOKING     = ["Non-smoker","Smoker","Ex-smoker","Shisha","Smoker + Shisha"]
REFERRAL    = [NA,"Self","Family","Physician","Psychologist","School","Other"]
HTYPE       = [NA,"Initial","Follow-up","Emergency","Consultation"]
ALIVE_M     = ["Alive","Deceased","Unknown"]
ALIVE_F     = ["Alive","Deceased","Unknown"]
CONS        = [NA,"No consanguinity","1st degree","2nd degree","3rd degree"]
PARENTS_REL = [NA,"Good","Fair","Poor","Separated","Divorced","One deceased"]
MARQ        = [NA,"Good","Fair","Poor","Separated"]
PRE_MAR     = [NA,"No prior relationship","Acquaintance only","Long relationship","Arranged","Other"]
NUM_CHILD   = [NA,"No children","1","2","3","4","5","6+"]
MAR_DUR     = [NA,"< 1 year","1–3 years","3–5 years","5–10 years","10+ years"]
ENGAGEMENT  = [NA,"No engagement","< 3 months","3–6 months","6–12 months","1+ year"]
ONSET_MODE  = [NA,"Sudden","Gradual"]
COURSE      = [NA,"Continuous","Recurrent episodes","Improving","Deteriorating","Fluctuating"]
COMPLIANCE  = [NA,"Compliant","Irregular","Non-compliant","Refusing"]
INSIGHT     = [NA,"Full","Partial","Absent"]
SLEEP       = ["Normal","Insomnia","Hypersomnia","Interrupted"]
APPETITE    = ["Normal","Decreased","Increased"]
SUICIDAL    = ["None","Passive ideation only","Active ideation","Clear plan"]
SUBSTANCE   = [NA,"None","Alcohol","Cannabis","Sedatives","Multiple","Other"]
HOBBIES     = ["Reading","Sports","Music","Art","Cooking","Gaming","Social media","None","Other"]
CHRONIC     = [NA,"None","Diabetes","Hypertension","Cardiac","Renal","Autoimmune","Cancer","Other"]
SIB_GENDER  = [NA,"Male","Female"]
SIB_EDU     = [NA,"Kindergarten","Primary","Preparatory","Secondary","University","Graduate","Not studying"]
BIRTH_ORDER = [NA,"1st","2nd","3rd","4th","5th","6th+","Only child"]
BIRTH_TYPE  = [NA,"Normal","Caesarean","Forceps","Vacuum"]
BIRTH_COMP  = [NA,"None","Jaundice","Incubator","Asphyxia","Low birth weight","Other"]
BF          = [NA,"Breastfed","Formula","Mixed"]
WEANING     = [NA,"< 6 months","6–12 months","12–18 months","18–24 months","After 2 years"]
MOTOR       = [NA,"Normal","Delayed","Early"]
SPEECH      = [NA,"Normal","Delayed","Absent","Regression after completion"]
TEETH       = [NA,"Normal (6–8 mo)","Early (< 6 mo)","Late (> 12 mo)"]
TOILET      = [NA,"Normal (18–30 mo)","Early","Late (> 3 years)"]
VACC        = [NA,"Complete","Incomplete","Unknown"]
ACADEMIC    = ["Excellent","Good","Average","Weak","Not studying"]
WANTED      = ["Yes, wanted","No, unwanted","Unplanned pregnancy"]
GENDER_DES  = ["Yes, desired gender","No, preferred other gender","No preference"]
LIVES_WITH  = [NA,"Both parents","Mother only","Father only","Grandparents","Relative","Other"]
SCREEN      = [NA,"< 1 hr","1–2 hrs","2–4 hrs","4–6 hrs","6+ hrs"]
PREG        = [NA,"Normal, no complications","Hypertension","Gestational diabetes","Bleeding","Advanced maternal age (>35)","Other complication"]
PUNISHMENT  = ["Verbal only","Privilege withdrawal","Physical","Ignoring","Mixed"]
STRESS_RX   = ["Calm","Crying","Aggression","Withdrawal","Mixed"]

# ══════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════

def sv(d, key, default="Not reported"):
    v = d.get(key, "")
    if not v: return default
    if isinstance(v, list): return ", ".join(v) if v else default
    v = str(v).strip()
    return v if v and v not in [NA, "—", "Not reported"] else default

def calc_age(birthdate_str):
    if not birthdate_str: return ""
    try:
        parts = re.split(r'[/\-\.]', birthdate_str.strip())
        if len(parts) == 3:
            d_, m_, y_ = int(parts[0]), int(parts[1]), int(parts[2])
            today = date.today()
            years  = today.year - y_ - ((today.month, today.day) < (m_, d_))
            months = (today.month - m_) % 12
            return f"{years} years, {months} months"
    except: pass
    return ""

def section(title_en, title_ar=""):
    ar_span = f'<span class="section-title-ar">/ {title_ar}</span>' if title_ar else ""
    st.markdown(f"""
    <div class="section-title">{title_en}{ar_span}</div>
    """, unsafe_allow_html=True)

def lbl(en, ar=""):
    ar_span = f'<span class="field-label-ar"> / {ar}</span>' if ar else ""
    st.markdown(f'<div class="field-label">{en}{ar_span}</div>', unsafe_allow_html=True)

def ti(en, ar, key, placeholder=""):
    lbl(en, ar)
    return st.text_input("", key=key, placeholder=placeholder, label_visibility="collapsed")

def ta(en, ar, key, height=100):
    lbl(en, ar)
    return st.text_area("", key=key, height=height, label_visibility="collapsed")

def rb(en, ar, opts, key, horizontal=True):
    lbl(en, ar)
    return st.radio("", opts, key=key, horizontal=horizontal, label_visibility="collapsed")

def sel(en, ar, opts, key):
    lbl(en, ar)
    return st.selectbox("", opts, key=key, label_visibility="collapsed")

def ms(en, ar, opts, key):
    lbl(en, ar)
    return st.multiselect("", opts, key=key, label_visibility="collapsed")

# ══════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════

with st.sidebar:
    if os.path.exists(LOGO_FILE):
        st.image(LOGO_FILE, use_container_width=True)
        st.markdown("---")

    st.markdown("## ⚙️ Settings")
    history_by = st.text_input("Specialist / Psychologist Name", placeholder="Dr. ...")
    st.markdown("---")
    st.markdown("## 📋 Form Type")
    sheet_type = st.radio(
        "",
        ["👤 Adult / بالغ", "👶 Child / طفل"],
        label_visibility="collapsed"
    )
    is_adult = "Adult" in sheet_type
    st.markdown("---")
    st.markdown(
        f'<div style="font-size:0.78rem;color:#8B7355;line-height:1.6;">'
        f'<b>{DOCTOR["name"]}</b><br>'
        f'{DOCTOR["title1"]}<br>'
        f'{DOCTOR["phone"]}</div>',
        unsafe_allow_html=True
    )

groq_key = st.secrets["GROQ_API_KEY"]

# ══════════════════════════════════════════════════════════════
#  PAGE HEADER
# ══════════════════════════════════════════════════════════════

st.markdown(f"""
<div class="page-header">
    <p>{'Adult Clinical History' if is_adult else 'Child Clinical History'} — {'بالغ' if is_adult else 'طفل'}</p>
    <h1>🧠 استمارة أخذ التاريخ المرضي</h1>
    <p>Dr. Hany Elhennawy Clinic — Neuro-Psychiatry</p>
</div>
""", unsafe_allow_html=True)

d = {}

# ══════════════════════════════════════════════════════════════
#  ══  ADULT FORM  ══
# ══════════════════════════════════════════════════════════════

if is_adult:

    # ── Personal Details ──────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Personal Details", "البيانات الشخصية")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["name"]      = ti("Full Name","الاسم الكامل","a_name")
        d["birthdate"] = ti("Date of Birth","تاريخ الميلاد","a_birthdate","DD/MM/YYYY")
        _age = calc_age(st.session_state.get("a_birthdate",""))
        d["age"] = _age
        if _age: st.markdown(f'<div class="age-badge">🎂 {_age}</div>', unsafe_allow_html=True)
        d["gender"]    = rb("Gender","النوع", GENDER,"a_gender")
    with c2:
        d["education"]  = sel("Education","المستوى التعليمي", EDU,"a_edu")
        d["occupation"] = sel("Occupation","الوظيفة", OCC,"a_occ")
        d["occ_detail"] = ti("Occupation details (if needed)","تفاصيل الوظيفة","a_occd")
        d["social"]     = sel("Social Status","الحالة الاجتماعية", SOCIAL,"a_social")
    with c3:
        d["smoking"]   = rb("Smoking","التدخين", SMOKING,"a_smoking", horizontal=False)
        d["referral"]  = sel("Referral Source","مصدر الإحالة", REFERRAL,"a_referral")
        d["htype"]     = sel("History Type","نوع التاريخ", HTYPE,"a_htype")
        d["phone"]     = ti("Phone","رقم الهاتف","a_phone")
        d["date"]      = ti("Session Date","تاريخ الجلسة","a_date", str(date.today()))
    lbl("Hobbies","الهوايات")
    d["hobbies"] = st.multiselect("", HOBBIES, key="a_hobbies", label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Family Details ────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Family Details", "بيانات الأسرة")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Father / الأب**")
        d["father_name"]  = ti("Father's Name","اسم الأب","a_fn")
        d["father_age"]   = ti("Father's Age","سن الأب","a_fa")
        d["father_occ"]   = ti("Father's Occupation","وظيفة الأب","a_fo")
        d["father_alive"] = rb("Father's Status","حالة الأب", ALIVE_M,"a_falive")
    with c2:
        st.markdown("**Mother / الأم**")
        d["mother_name"]  = ti("Mother's Name","اسم الأم","a_mn")
        d["mother_age"]   = ti("Mother's Age","سن الأم","a_ma")
        d["mother_occ"]   = ti("Mother's Occupation","وظيفة الأم","a_mo")
        d["mother_alive"] = rb("Mother's Status","حالة الأم", ALIVE_F,"a_malive")
    c1, c2, c3 = st.columns(3)
    with c1: d["consanguinity"]    = sel("Consanguinity","القرابة بين الأبوين", CONS,"a_cons")
    with c2: d["parents_together"] = rb("Parents living together?","هل يعيشان معاً؟", YES_NO_NA,"a_ptog")
    with c3: d["chronic"]          = sel("Chronic illness in family","مرض مزمن في الأسرة", CHRONIC,"a_chronic")
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Marriage Details ──────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Marriage Details", "بيانات الزواج")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["spouse_name"]  = ti("Spouse Name","اسم الزوج/ة","a_spn")
        d["spouse_age"]   = ti("Spouse Age","سن الزوج/ة","a_spa")
        d["spouse_occ"]   = sel("Spouse Occupation","وظيفة الزوج/ة", OCC,"a_spo")
    with c2:
        d["marriage_dur"] = sel("Marriage Duration","مدة الزواج", MAR_DUR,"a_mdur")
        d["engagement"]   = sel("Engagement Period","فترة الخطوبة", ENGAGEMENT,"a_eng")
        d["num_children"] = sel("Number of Children","عدد الأبناء", NUM_CHILD,"a_nch")
    with c3:
        d["katb"]         = rb("Katb Ketab before marriage?","كتب كتاب قبل الزواج؟",["Yes","No","N/A"],"a_katb")
        d["marriage_qual"]= sel("Marriage Quality","جودة الزواج", MARQ,"a_mqual")
        d["pre_marriage"] = sel("Relationship before marriage","العلاقة قبل الزواج", PRE_MAR,"a_pre")
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Siblings ──────────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Siblings", "الإخوة والأخوات")
    siblings = []
    for i in range(1, 5):
        c1,c2,c3,c4,c5 = st.columns([1,2,1,2,2])
        with c1:
            lbl(f"Gender {i}","")
            g = st.selectbox("",SIB_GENDER,key=f"a_sg{i}",label_visibility="collapsed")
        with c2:
            n = st.text_input("",key=f"a_sn{i}",placeholder=f"Name {i}",label_visibility="collapsed")
        with c3:
            a_ = st.text_input("",key=f"a_sa{i}",placeholder=f"Age {i}",label_visibility="collapsed")
        with c4:
            lbl(f"Edu {i}","")
            e = st.selectbox("",SIB_EDU,key=f"a_se{i}",label_visibility="collapsed")
        with c5:
            nt = st.text_input("",key=f"a_st{i}",placeholder=f"Notes {i}",label_visibility="collapsed")
        if n: siblings.append({"gender":g,"name":n,"age":a_,"edu":e,"notes":nt})
    d["siblings"] = siblings
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Complaints & HPI ──────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Complaints & History of Present Illness", "الشكاوى والتاريخ المرضي الحالي")
    c1,c2,c3 = st.columns(3)
    with c1: d["onset"]      = ti("Symptom onset","متى بدأت الأعراض؟","a_onset")
    with c2: d["onset_mode"] = sel("Mode of onset","طريقة البداية", ONSET_MODE,"a_omode")
    with c3: d["course"]     = sel("Course","مسار المرض", COURSE,"a_course")
    d["complaints"] = ta("Chief Complaints (C/O)","الشكاوى الرئيسية","a_co",120)
    d["hpi"]        = ta("History of Present Illness (HPI)","تاريخ المرض الحالي بالتفصيل","a_hpi",200)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Drug History ──────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Drug History", "تاريخ الأدوية")
    c1,c2 = st.columns(2)
    with c1: d["on_meds"]    = rb("Currently on medication?","يتناول أدوية حالياً؟", YES_NO_NA,"a_onmeds")
    with c2: d["compliance"] = sel("Compliance","الالتزام بالأدوية", COMPLIANCE,"a_comp")
    d["drug_hx"] = ta("Medications (name, dose, duration)","تفاصيل الأدوية","a_drug",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Past History ──────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Past Medical & Psychiatric History", "التاريخ المرضي السابق")
    c1,c2 = st.columns(2)
    with c1: d["prev_psych"] = rb("Previous psychiatric illness?","مرض نفسي سابق؟", YES_NO_NA,"a_ppsych")
    with c2: d["prev_hosp"]  = rb("Previous hospitalization?","دخول مستشفى سابق؟", YES_NO_NA,"a_phosp")
    d["past_hx"] = ta("Details","تفاصيل","a_past",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Family History ────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Family Psychiatric & Neurological History", "التاريخ العائلي")
    c1,c2 = st.columns(2)
    with c1:
        d["fam_psych"] = rb("Psychiatric illness in family?","مرض نفسي في الأسرة؟", YES_NO_NA,"a_fpsych")
        if st.session_state.get("a_fpsych","").startswith("Yes"):
            d["fam_psych_details"] = ti("Specify illness & family member","التفاصيل","a_fpsych_det")
        else: d["fam_psych_details"] = ""
    with c2:
        d["fam_neuro"] = rb("Neurological illness in family?","مرض عصبي في الأسرة؟", YES_NO_NA,"a_fneuro")
        if st.session_state.get("a_fneuro","").startswith("Yes"):
            d["fam_neuro_details"] = ti("Specify illness & family member","التفاصيل","a_fneuro_det")
        else: d["fam_neuro_details"] = ""
    d["family_hx"] = ta("Details","تفاصيل","a_famhx",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Investigations ────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Investigations", "الفحوصات")
    d["had_inv"]        = rb("Investigations done?","أُجريت فحوصات؟", YES_NO_NA,"a_hadinv")
    d["investigations"] = ta("Details & results","التفاصيل والنتائج","a_inv",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Surgeries ─────────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Operations & Surgeries", "العمليات والجراحات")
    d["had_surg"]  = rb("Previous surgeries?","عمليات سابقة؟", YES_NO_NA,"a_hsurg")
    d["surgeries"] = ta("Details","تفاصيل","a_surg",60)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Clinical Assessment ───────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Clinical Assessment", "التقييم السريري")
    c1,c2 = st.columns(2)
    with c1:
        d["sleep"]    = rb("Sleep pattern","نمط النوم", SLEEP,"a_sleep")
        d["appetite"] = rb("Appetite","الشهية", APPETITE,"a_appetite")
        d["insight"]  = sel("Insight / Awareness","البصيرة", INSIGHT,"a_insight")
    with c2:
        d["suicidal"]  = rb("Suicidal ideation","أفكار انتحارية", SUICIDAL,"a_suicidal", horizontal=False)
        d["substance"] = sel("Substance use","تعاطي مواد", SUBSTANCE,"a_subs")
        d["substance_details"] = ta("Substance details","تفاصيل","a_subsd",60)
    d["extra_notes"] = ta("Additional notes","ملاحظات إضافية","a_extra",80)
    st.markdown('</div>', unsafe_allow_html=True)

    patient_name = d.get("name") or "Patient"

# ══════════════════════════════════════════════════════════════
#  ══  CHILD FORM  ══
# ══════════════════════════════════════════════════════════════

else:

    # ── Personal Details ──────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Child's Personal Details", "البيانات الشخصية للطفل")
    c1,c2,c3 = st.columns(3)
    with c1:
        d["name"]       = ti("Child's Full Name","اسم الطفل كاملاً","c_name")
        d["birthdate"]  = ti("Date of Birth","تاريخ الميلاد","c_birthdate","DD/MM/YYYY")
        _age2 = calc_age(st.session_state.get("c_birthdate",""))
        d["age"] = _age2
        if _age2: st.markdown(f'<div class="age-badge">🎂 {_age2}</div>', unsafe_allow_html=True)
        d["gender"]      = rb("Gender","النوع", GENDER,"c_gender")
        d["birth_order"] = sel("Birth Order","ترتيب الميلاد", BIRTH_ORDER,"c_border")
    with c2:
        d["lives_with"]  = sel("Lives with","يعيش مع", LIVES_WITH,"c_lives")
        d["school"]      = ti("School Name","اسم المدرسة","c_school")
        d["grade"]       = ti("Grade / Class","الصف الدراسي","c_grade")
        d["academic"]    = rb("Academic Performance","المستوى الدراسي", ACADEMIC,"c_academic")
        d["screen_time"] = sel("Daily Screen Time","وقت الشاشة اليومي", SCREEN,"c_screen")
    with c3:
        d["phone"]      = ti("Phone","تليفون","c_phone")
        d["date"]       = ti("Session Date","تاريخ الجلسة","c_date", str(date.today()))
        d["wanted"]     = rb("Was the child wanted?","هل كان مرغوباً فيه؟", WANTED,"c_wanted", horizontal=False)
        d["gender_des"] = rb("Was the gender desired?","هل كان النوع مرغوباً؟", GENDER_DES,"c_gdes", horizontal=False)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Developmental Milestones ──────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Developmental History", "مراحل النمو")
    c1,c2,c3 = st.columns(3)
    with c1:
        st.markdown("**Pregnancy & Birth / الحمل والولادة**")
        d["pregnancy"]   = ta("Pregnancy details","تفاصيل الحمل","c_preg",80)
        d["birth_type"]  = sel("Birth type","نوع الولادة", BIRTH_TYPE,"c_btype")
        d["birth_comp"]  = sel("Birth complications","مضاعفات الولادة", BIRTH_COMP,"c_bcomp")
        d["vacc_status"] = sel("Vaccinations","التطعيمات", VACC,"c_vacc")
        d["vacc_comp"]   = ti("Post-vaccine complications","مضاعفات التطعيم","c_vcomp")
    with c2:
        st.markdown("**Feeding & Growth / التغذية والنمو**")
        d["breastfeeding"]= sel("Breastfeeding","الرضاعة", BF,"c_bf")
        d["weaning"]      = sel("Weaning age","سن الفطام", WEANING,"c_wean")
        d["motor"]        = sel("Motor development","النمو الحركي", MOTOR,"c_motor")
        d["motor_detail"] = ti("Motor details (walk, sit...)","تفاصيل الحركة","c_motord")
        d["teething"]     = sel("Teething","التسنين", TEETH,"c_teeth")
        d["toilet"]       = sel("Toilet training","تدريب دورة المياه", TOILET,"c_toilet")
    with c3:
        st.markdown("**Language & Cognition / اللغة والإدراك**")
        d["speech"]        = sel("Speech","الكلام", SPEECH,"c_speech")
        d["speech_detail"] = ti("Speech details","تفاصيل الكلام","c_speechd")
        d["attention"]     = rb("Attention","الانتباه",["Normal","Weak","N/A"],"c_attn")
        d["concentration"] = rb("Concentration","التركيز",["Normal","Weak","N/A"],"c_conc")
        d["comprehension"] = rb("Comprehension","الفهم والإدراك",["Normal","Weak","N/A"],"c_comp")
    d["dev_notes"] = ta("Developmental notes","ملاحظات النمو","c_devnotes",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Family Details ────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Family Details", "بيانات الأسرة")
    c1,c2 = st.columns(2)
    with c1:
        st.markdown("**Father / الأب**")
        d["father_name"]       = ti("Father's Name","اسم الأب","c_fn")
        d["father_age"]        = ti("Father's Age","سن الأب","c_fa")
        d["father_occ"]        = ti("Father's Occupation","وظيفة الأب","c_fo")
        d["father_alive"]      = rb("Father's Status","حالة الأب", ALIVE_M,"c_falive")
        d["father_hereditary"] = ti("Father — hereditary illness (if any)","مرض وراثي عند الأب","c_fh")
    with c2:
        st.markdown("**Mother / الأم**")
        d["mother_name"]       = ti("Mother's Name","اسم الأم","c_mn")
        d["mother_age"]        = ti("Mother's Age","سن الأم","c_ma")
        d["mother_occ"]        = ti("Mother's Occupation","وظيفة الأم","c_mo")
        d["mother_alive"]      = rb("Mother's Status","حالة الأم", ALIVE_F,"c_malive")
        d["mother_hereditary"] = ti("Mother — hereditary illness (if any)","مرض وراثي عند الأم","c_mh")
    c1,c2 = st.columns(2)
    with c1: d["consanguinity"] = sel("Consanguinity","القرابة بين الأبوين", CONS,"c_cons")
    with c2: d["parents_rel"]   = sel("Parents relationship","طبيعة العلاقة الزوجية", PARENTS_REL,"c_prel")
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Siblings ──────────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Siblings", "الإخوة والأخوات")
    siblings = []
    for i in range(1,5):
        c1,c2,c3,c4,c5 = st.columns([1,2,1,2,2])
        with c1:
            lbl(f"Gender {i}","")
            g = st.selectbox("",SIB_GENDER,key=f"c_sg{i}",label_visibility="collapsed")
        with c2:
            n = st.text_input("",key=f"c_sn{i}",placeholder=f"Name {i}",label_visibility="collapsed")
        with c3:
            a_ = st.text_input("",key=f"c_sa{i}",placeholder=f"Age {i}",label_visibility="collapsed")
        with c4:
            lbl(f"Edu {i}","")
            e = st.selectbox("",SIB_EDU,key=f"c_se{i}",label_visibility="collapsed")
        with c5:
            nt = st.text_input("",key=f"c_st{i}",placeholder=f"Notes {i}",label_visibility="collapsed")
        if n: siblings.append({"gender":g,"name":n,"age":a_,"edu":e,"notes":nt})
    d["siblings"] = siblings
    c1,c2 = st.columns(2)
    with c1:
        lbl("Sibling relationship","علاقة الأخوة ببعض")
        d["sibling_rel"] = st.selectbox("", [NA,"Good","Fair","Competitive","Constant conflict","Mutual neglect"], key="c_sibrel", label_visibility="collapsed")
    with c2:
        d["same_school"] = rb("Siblings at same school?","في نفس المدرسة؟",["Yes","No","N/A"],"c_ssch")
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Complaints & HPI ──────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Complaints & History of Present Illness", "الشكاوى والتاريخ المرضي الحالي")
    c1,c2,c3 = st.columns(3)
    with c1: d["onset"]      = ti("Symptom onset","متى بدأت الأعراض؟","c_onset")
    with c2: d["onset_mode"] = sel("Mode of onset","طريقة البداية", ONSET_MODE,"c_omode")
    with c3: d["course"]     = sel("Course","مسار المرض", COURSE,"c_course")
    d["complaints"] = ta("Chief Complaints (C/O)","الشكاوى الرئيسية","c_co",120)
    d["hpi"]        = ta("History of Present Illness (HPI)","تاريخ المرض الحالي بالتفصيل","c_hpi",200)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Past History ──────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Past Medical History", "التاريخ المرضي السابق")
    c1,c2,c3 = st.columns(3)
    with c1:
        d["high_fever"]  = rb("High fever ≥40°C?","حرارة ≥40 درجة؟", YES_NO_NA,"c_hfever")
        d["head_trauma"] = rb("Head trauma?","ارتطام رأس؟", YES_NO_NA,"c_htrauma")
        if st.session_state.get("c_htrauma","").startswith("Yes"):
            d["head_trauma_location"] = ti("Location on head","مكان الارتطام","c_htrauma_loc")
            d["head_trauma_details"]  = ti("How did it happen?","كيف حدث؟","c_htrauma_det")
        else:
            d["head_trauma_location"] = ""; d["head_trauma_details"] = ""
    with c2:
        d["convulsions"]  = rb("Convulsions?","تشنجات؟", YES_NO_NA,"c_conv")
        d["post_vaccine"] = rb("Post-vaccine complications?","مضاعفات بعد التطعيم؟", YES_NO_NA,"c_pvacc")
    with c3:
        d["prev_hosp"]    = rb("Previous hospitalization?","دخول مستشفى سابق؟", YES_NO_NA,"c_phosp")
        d["prev_therapy"] = rb("Previous therapy sessions?","جلسات علاجية سابقة؟", YES_NO_NA,"c_pther")
    d["past_hx"] = ta("Details","تفاصيل","c_past",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Family History ────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Family History", "التاريخ العائلي")
    c1,c2 = st.columns(2)
    with c1:
        d["fam_psych"] = rb("Psychiatric illness in family?","مرض نفسي في الأسرة؟", YES_NO_NA,"c_fpsych")
        if st.session_state.get("c_fpsych","").startswith("Yes"):
            d["fam_psych_details"] = ti("Specify","التفاصيل","c_fpsych_det")
        else: d["fam_psych_details"] = ""
        d["fam_neuro"] = rb("Neurological illness in family?","مرض عصبي في الأسرة؟", YES_NO_NA,"c_fneuro")
        if st.session_state.get("c_fneuro","").startswith("Yes"):
            d["fam_neuro_details"] = ti("Specify","التفاصيل","c_fneuro_det")
        else: d["fam_neuro_details"] = ""
    with c2:
        d["fam_mr"] = rb("Intellectual disability in family?","إعاقة ذهنية في الأسرة؟", YES_NO_NA,"c_fmr")
        if st.session_state.get("c_fmr","").startswith("Yes"):
            d["fam_mr_details"] = ti("Specify","التفاصيل","c_fmr_det")
        else: d["fam_mr_details"] = ""
        d["fam_epilepsy"] = rb("Epilepsy in family?","صرع في الأسرة؟", YES_NO_NA,"c_fepil")
        if st.session_state.get("c_fepil","").startswith("Yes"):
            d["fam_epilepsy_details"] = ti("Specify","التفاصيل","c_fepil_det")
        else: d["fam_epilepsy_details"] = ""
    d["family_hx"] = ta("Details","تفاصيل","c_famhx",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Investigations ────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Investigations", "الفحوصات")
    c1,c2,c3 = st.columns(3)
    with c1:
        d["had_ct"]  = rb("CT scan?","أشعة مقطعية؟", YES_NO_NA,"c_ct")
        d["had_mri"] = rb("MRI?","رنين مغناطيسي؟", YES_NO_NA,"c_mri")
    with c2:
        d["had_eeg"] = rb("EEG?","رسم مخ؟", YES_NO_NA,"c_eeg")
        d["had_iq"]  = rb("IQ test (SB5)?","اختبار ذكاء SB5؟", YES_NO_NA,"c_iq")
    with c3:
        d["had_cars"]   = rb("CARS scale?","مقياس CARS؟", YES_NO_NA,"c_cars")
        d["cars_score"] = ti("CARS score (if done)","درجة CARS","c_carsscore")
    d["investigations"] = ta("Details & results","التفاصيل والنتائج","c_inv",80)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Surgeries ─────────────────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Operations & Surgeries", "العمليات والجراحات")
    d["had_surg"]  = rb("Previous surgeries?","عمليات سابقة؟", YES_NO_NA,"c_hsurg")
    d["surgeries"] = ta("Details","تفاصيل","c_surg",60)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Clinical Assessment ───────────────────────────────────
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    section("Clinical Assessment", "التقييم السريري")
    c1,c2 = st.columns(2)
    with c1:
        d["sleep"]    = rb("Sleep pattern","نمط النوم", SLEEP,"c_sleep")
        d["appetite"] = rb("Appetite","الشهية", APPETITE,"c_appetite")
        lbl("Punishment methods used","طرق العقاب المستخدمة")
        d["punishment"] = st.multiselect("",["Verbal only","Privilege withdrawal","Physical","Ignoring","Mixed"],key="c_punish",label_visibility="collapsed")
    with c2:
        lbl("Reaction to stress","رد الفعل تجاه الضغوط")
        d["stress_reaction"] = st.multiselect("",["Calm","Crying","Aggression","Withdrawal","Anger episodes","Enuresis","Other"],key="c_stress",label_visibility="collapsed")
        d["therapy"] = ta("Current therapy sessions","الجلسات العلاجية الحالية","c_therapy",80)
    d["extra_notes"] = ta("Additional notes","ملاحظات إضافية","c_extra",80)
    st.markdown('</div>', unsafe_allow_html=True)

    patient_name = d.get("name") or "Child"

# ══════════════════════════════════════════════════════════════
#  GENERATE REPORT BUTTON
# ══════════════════════════════════════════════════════════════

st.markdown("---")
col_btn = st.columns([1,2,1])
with col_btn[1]:
    generate = st.button("✦ Generate Report / توليد التقرير", type="primary", use_container_width=True)

# ══════════════════════════════════════════════════════════════
#  VERBATIM ARABIC EXTRACTION (bypasses AI — appended directly)
# ══════════════════════════════════════════════════════════════

def extract_arabic_verbatim(d, is_adult):
    """Extract all free-text Arabic fields directly — never passed through AI."""
    fields = []
    if is_adult:
        pairs = [
            ("الشكوى الرئيسية (C/O)", "complaints"),
            ("تاريخ المرض الحالي (HPI)", "hpi"),
            ("تفاصيل الأدوية", "drug_hx"),
            ("التاريخ المرضي السابق — تفاصيل", "past_hx"),
            ("التاريخ العائلي — تفاصيل", "family_hx"),
            ("الفحوصات — تفاصيل", "investigations"),
            ("العمليات — تفاصيل", "surgeries"),
            ("تفاصيل تعاطي المواد", "substance_details"),
            ("ملاحظات إضافية", "extra_notes"),
        ]
    else:
        pairs = [
            ("الشكوى الرئيسية (C/O)", "complaints"),
            ("تاريخ المرض الحالي (HPI)", "hpi"),
            ("تفاصيل الحمل", "pregnancy"),
            ("ملاحظات النمو", "dev_notes"),
            ("التاريخ المرضي السابق — تفاصيل", "past_hx"),
            ("التاريخ العائلي — تفاصيل", "family_hx"),
            ("الفحوصات — تفاصيل", "investigations"),
            ("العمليات — تفاصيل", "surgeries"),
            ("الجلسات العلاجية الحالية", "therapy"),
            ("ملاحظات إضافية", "extra_notes"),
        ]
    for label, key in pairs:
        val = d.get(key, "").strip() if isinstance(d.get(key,""), str) else ""
        if val and val not in ["Not reported","لم يُذكر",""]:
            fields.append((label, val))
    return fields

# ══════════════════════════════════════════════════════════════
#  DATA BLOCK BUILDER
# ══════════════════════════════════════════════════════════════

def build_data_block(d, is_adult, history_by):
    sib_text = "\n".join([
        f"  {i+1}. {sb['name']} | {sb['gender']} | Age: {sb['age']} | Edu: {sb['edu']} | Notes: {sb['notes'] or 'None'}"
        for i, sb in enumerate(d.get("siblings",[]))
    ]) or "No siblings entered"

    if is_adult:
        return f"""
Patient: {sv(d,'name')} | DOB: {sv(d,'birthdate')} | Age: {sv(d,'age')} | Gender: {sv(d,'gender')}
Session date: {sv(d,'date')} | Specialist: {history_by or 'Not reported'} | History type: {sv(d,'htype')}
Phone: {sv(d,'phone')} | Referral: {sv(d,'referral')}
Education: {sv(d,'education')} | Occupation: {sv(d,'occupation')} — {sv(d,'occ_detail')}
Social status: {sv(d,'social')} | Smoking: {sv(d,'smoking')} | Hobbies: {sv(d,'hobbies')}

FAMILY:
Father: {sv(d,'father_name')} | Age: {sv(d,'father_age')} | Occ: {sv(d,'father_occ')} | Status: {sv(d,'father_alive')}
Mother: {sv(d,'mother_name')} | Age: {sv(d,'mother_age')} | Occ: {sv(d,'mother_occ')} | Status: {sv(d,'mother_alive')}
Consanguinity: {sv(d,'consanguinity')} | Parents together: {sv(d,'parents_together')} | Chronic illness: {sv(d,'chronic')}

MARRIAGE:
Spouse: {sv(d,'spouse_name')} | Age: {sv(d,'spouse_age')} | Occ: {sv(d,'spouse_occ')}
Duration: {sv(d,'marriage_dur')} | Engagement: {sv(d,'engagement')} | Children: {sv(d,'num_children')}
Katb Ketab: {sv(d,'katb')} | Quality: {sv(d,'marriage_qual')} | Pre-marriage: {sv(d,'pre_marriage')}

SIBLINGS:
{sib_text}

PRESENTING COMPLAINT:
Onset: {sv(d,'onset')} | Mode: {sv(d,'onset_mode')} | Course: {sv(d,'course')}
[Chief complaints and HPI provided in Arabic — see Original Arabic Responses section]

DRUG HISTORY:
On medication: {sv(d,'on_meds')} | Compliance: {sv(d,'compliance')}
[Drug details in Arabic — see Original Arabic Responses section]

PAST HISTORY:
Previous psychiatric: {sv(d,'prev_psych')} | Previous hospitalization: {sv(d,'prev_hosp')}
[Details in Arabic — see Original Arabic Responses section]

FAMILY HISTORY:
Psychiatric: {sv(d,'fam_psych')}{(' — '+sv(d,'fam_psych_details')) if d.get('fam_psych_details') else ''} | Neurological: {sv(d,'fam_neuro')}{(' — '+sv(d,'fam_neuro_details')) if d.get('fam_neuro_details') else ''}

INVESTIGATIONS: Done: {sv(d,'had_inv')}
SURGERIES: Done: {sv(d,'had_surg')}

CLINICAL ASSESSMENT:
Sleep: {sv(d,'sleep')} | Appetite: {sv(d,'appetite')} | Suicidal ideation: {sv(d,'suicidal')} | Insight: {sv(d,'insight')}
Substance use: {sv(d,'substance')}
"""
    else:
        return f"""
Child: {sv(d,'name')} | DOB: {sv(d,'birthdate')} | Age: {sv(d,'age')} | Gender: {sv(d,'gender')}
Session date: {sv(d,'date')} | Specialist: {history_by or 'Not reported'}
Phone: {sv(d,'phone')} | Lives with: {sv(d,'lives_with')}
School: {sv(d,'school')} | Grade: {sv(d,'grade')} | Academic: {sv(d,'academic')}
Birth order: {sv(d,'birth_order')} | Screen time: {sv(d,'screen_time')}
Child wanted: {sv(d,'wanted')} | Gender desired: {sv(d,'gender_des')}

DEVELOPMENTAL HISTORY:
Birth type: {sv(d,'birth_type')} | Birth complications: {sv(d,'birth_comp')}
Vaccinations: {sv(d,'vacc_status')} | Post-vaccine complications: {sv(d,'vacc_comp')}
Breastfeeding: {sv(d,'breastfeeding')} | Weaning: {sv(d,'weaning')}
Motor: {sv(d,'motor')} — {sv(d,'motor_detail')}
Teething: {sv(d,'teething')} | Toilet training: {sv(d,'toilet')}
Speech: {sv(d,'speech')} — {sv(d,'speech_detail')}
Attention: {sv(d,'attention')} | Concentration: {sv(d,'concentration')} | Comprehension: {sv(d,'comprehension')}

FAMILY:
Father: {sv(d,'father_name')} | Age: {sv(d,'father_age')} | Occ: {sv(d,'father_occ')} | Status: {sv(d,'father_alive')} | Hereditary: {sv(d,'father_hereditary')}
Mother: {sv(d,'mother_name')} | Age: {sv(d,'mother_age')} | Occ: {sv(d,'mother_occ')} | Status: {sv(d,'mother_alive')} | Hereditary: {sv(d,'mother_hereditary')}
Consanguinity: {sv(d,'consanguinity')} | Parents relationship: {sv(d,'parents_rel')}

SIBLINGS:
{sib_text}
Sibling relationship: {sv(d,'sibling_rel')} | Same school: {sv(d,'same_school')}

PRESENTING COMPLAINT:
Onset: {sv(d,'onset')} | Mode: {sv(d,'onset_mode')} | Course: {sv(d,'course')}
[Chief complaints and HPI provided in Arabic — see Original Arabic Responses section]

PAST HISTORY:
High fever ≥40°C: {sv(d,'high_fever')} | Head trauma: {sv(d,'head_trauma')}{(' — Location: '+sv(d,'head_trauma_location')+' — Details: '+sv(d,'head_trauma_details')) if d.get('head_trauma_location') else ''}
Convulsions: {sv(d,'convulsions')} | Post-vaccine complications: {sv(d,'post_vaccine')}
Previous hospitalization: {sv(d,'prev_hosp')} | Previous therapy: {sv(d,'prev_therapy')}

FAMILY HISTORY:
Psychiatric: {sv(d,'fam_psych')}{(' — '+sv(d,'fam_psych_details')) if d.get('fam_psych_details') else ''} | Neurological: {sv(d,'fam_neuro')}{(' — '+sv(d,'fam_neuro_details')) if d.get('fam_neuro_details') else ''}
Intellectual disability: {sv(d,'fam_mr')}{(' — '+sv(d,'fam_mr_details')) if d.get('fam_mr_details') else ''} | Epilepsy: {sv(d,'fam_epilepsy')}{(' — '+sv(d,'fam_epilepsy_details')) if d.get('fam_epilepsy_details') else ''}

INVESTIGATIONS: CT: {sv(d,'had_ct')} | MRI: {sv(d,'had_mri')} | EEG: {sv(d,'had_eeg')} | SB5: {sv(d,'had_iq')} | CARS: {sv(d,'had_cars')} — Score: {sv(d,'cars_score')}
SURGERIES: Done: {sv(d,'had_surg')}

CLINICAL ASSESSMENT:
Sleep: {sv(d,'sleep')} | Appetite: {sv(d,'appetite')}
Punishment methods: {sv(d,'punishment')} | Stress reaction: {sv(d,'stress_reaction')}
"""

# ══════════════════════════════════════════════════════════════
#  GROQ PROMPT
# ══════════════════════════════════════════════════════════════

def build_prompt(data_block, is_adult, history_by):
    return f"""You are a clinical report formatter for a neuro-psychiatry clinic.
Generate a COMPACT, professional clinical history report in ENGLISH ONLY.
TARGET: 2–3 pages maximum.

STRICT RULES:
1. Write ALL sections in English only. Zero Arabic in your output.
2. DO NOT translate, summarize, or paraphrase any Arabic text — leave that entirely to the "Original Arabic Responses" section which will be appended separately after your output. Simply write "[Arabic responses appended below]" as a placeholder for that section.
3. DO NOT add diagnosis, interpretation, clinical judgment, or assumptions.
4. DO NOT invent or add any information not present in the input data.
5. Skip any field marked "Not reported". Skip empty sections entirely.
6. "No" answers: omit unless clinically significant.
7. NO markdown symbols (no **, no ##, no ---, no bullet dashes). 
8. Section titles: ALL CAPS numbered. Example: 1. PATIENT INFORMATION
9. Use pipe-format tables: Field | Value
10. Group related short fields inline: DOB: [v]  |  Age: [v]  |  Gender: [v]
11. Inline observations (no table): Sleep: Normal  |  Appetite: Decreased

REPORT STRUCTURE:

REPORT HEADER:
Patient Name | [value]
Form Type | {"Adult" if is_adult else "Child"}
Specialist | {history_by or "Not reported"}
Date | [value]  |  Phone | [value]
━━━━━━━━━━━━━━━━━━━━━━━━━━━

CLINICAL SUMMARY
2–3 concise clinical sentences. English only. No Arabic. No diagnosis.

1. PATIENT INFORMATION
Compact table: all personal/social/lifestyle fields.

2. PRESENTING CONCERNS
Part A — small table: Onset | Mode | Course
Part B — list the symptoms/complaints in plain English lines (translate from structured fields only, not from free-text Arabic).

3. {"FAMILY & MARRIAGE BACKGROUND" if is_adult else "FAMILY BACKGROUND"}
Parents table + {"marriage table + " if is_adult else ""}consanguinity + siblings table.

{"4. DEVELOPMENTAL HISTORY" if not is_adult else ""}
{"Compact tables: Birth/Feeding table + Development table." if not is_adult else ""}

{"4. PAST HISTORY" if is_adult else "5. PAST HISTORY"}
Compact table.

{"5. FAMILY HISTORY" if is_adult else "6. FAMILY HISTORY"}
Compact table.

{"6. MEDICAL OBSERVATIONS" if is_adult else "7. MEDICAL OBSERVATIONS"}
Inline format per line.

{"7. INVESTIGATIONS" if is_adult else "8. INVESTIGATIONS"}
Compact table.

{"8. ORIGINAL ARABIC RESPONSES" if is_adult else "9. ORIGINAL ARABIC RESPONSES"}
Write exactly this line and nothing else:
[Arabic responses appended below]

DATA:
{data_block}
Specialist: {history_by or "Not reported"} | Form: {"Adult" if is_adult else "Child"}
"""

# ══════════════════════════════════════════════════════════════
#  REPORT GENERATION
# ══════════════════════════════════════════════════════════════

if generate:
    data_block = build_data_block(d, is_adult, history_by)
    prompt = build_prompt(data_block, is_adult, history_by)

    with st.spinner("Generating report / جاري إنشاء التقرير..."):
        try:
            client = Groq(api_key=groq_key)
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role":"user","content":prompt}],
                max_tokens=4000,
                temperature=0.3,
            )
            ai_report = response.choices[0].message.content

            # ── Append Arabic verbatim DIRECTLY (bypasses AI) ──
            arabic_fields = extract_arabic_verbatim(d, is_adult)
            arabic_section = ""
            if arabic_fields:
                arabic_section = "\n\nالردود العربية الأصلية / ORIGINAL ARABIC RESPONSES\n"
                arabic_section += "=" * 50 + "\n"
                for label, text in arabic_fields:
                    arabic_section += f"\n{label}:\n{text}\n"

            # Replace placeholder with actual Arabic content
            full_report = ai_report.replace("[Arabic responses appended below]", arabic_section.strip())
            # In case AI didn't use the placeholder
            if arabic_section and "[Arabic responses appended below]" not in ai_report:
                full_report = ai_report + "\n" + arabic_section

            st.session_state["report_text"]  = full_report
            st.session_state["report_pname"] = patient_name
            st.session_state["report_sheet"] = "Adult" if is_adult else "Child"
            st.session_state["report_by"]    = history_by or "—"
            st.session_state["report_d"]     = d
        except Exception as e:
            st.error(f"Error generating report: {str(e)}")

# ══════════════════════════════════════════════════════════════
#  DISPLAY & DOWNLOAD REPORT
# ══════════════════════════════════════════════════════════════

if st.session_state.get("report_text"):
    rt  = st.session_state["report_text"]
    pn  = st.session_state.get("report_pname","Patient")
    rs  = st.session_state.get("report_sheet","")
    rb_ = st.session_state.get("report_by","—")
    fn_base = pn.replace(' ','_')

    st.markdown("---")
    st.markdown(f"""
    <div class="info-box">
        ✅ Report generated for <strong>{pn}</strong> — {rs} form
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="report-box">', unsafe_allow_html=True)
    st.text_area("", value=rt, height=600, label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)

    # ══ BUILD DOCX ══════════════════════════════════════════
    def build_docx(rt, pn, rs, rb_):
        doc = Document()
        for section_ in doc.sections:
            section_.top_margin    = Cm(1.8); section_.bottom_margin = Cm(1.8)
            section_.left_margin   = Cm(2.0); section_.right_margin  = Cm(2.0)
            section_.different_first_page_header_footer = True
            for hdr in [section_.header, section_.first_page_header]:
                for p in hdr.paragraphs: p.clear()

        # Page border
        for section_ in doc.sections:
            sectPr = section_._sectPr
            pgB = OxmlElement('w:pgBorders')
            pgB.set(qn('w:offsetFrom'),'page')
            for side in ('top','left','bottom','right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'12')
                b.set(qn('w:space'),'24'); b.set(qn('w:color'),'1A5CB8')
                pgB.append(b)
            sectPr.append(pgB)

        # Footer page numbers
        for section_ in doc.sections:
            footer = section_.footer
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.clear(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(); run.font.size = Pt(9); run.font.color.rgb = CLINIC_BLUE
            for tag, text in [('begin',None),(None,' PAGE '),('end',None)]:
                if tag:
                    el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),tag); run._r.append(el)
                else:
                    instr = OxmlElement('w:instrText'); instr.text = text; run._r.append(instr)

        # Header with logo
        p_top = doc.add_paragraph()
        p_top.paragraph_format.space_before = Pt(0)
        p_top.paragraph_format.space_after  = Pt(4)
        if os.path.exists(LOGO_FILE):
            p_top.add_run().add_picture(LOGO_FILE, width=Inches(1.2))
        r_t = p_top.add_run("   Clinical History Report")
        r_t.font.name="Arial"; r_t.font.size=Pt(18)
        r_t.font.bold=True; r_t.font.color.rgb=CLINIC_BLUE
        pPr = p_top._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
        bot.set(qn('w:space'),'4');    bot.set(qn('w:color'),'1A5CB8')
        pBdr.append(bot); pPr.append(pBdr)
        doc.add_paragraph()

        def add_section_title(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(text.strip())
            r.font.size=Pt(13); r.font.name="Arial"
            r.font.bold=True; r.font.color.rgb=CLINIC_BLUE
            pPr2 = p._p.get_or_add_pPr()
            pBdr2 = OxmlElement('w:pBdr')
            bot2 = OxmlElement('w:bottom')
            bot2.set(qn('w:val'),'single'); bot2.set(qn('w:sz'),'6')
            bot2.set(qn('w:space'),'2');    bot2.set(qn('w:color'),'1A5CB8')
            pBdr2.append(bot2); pPr2.append(pBdr2)

        def add_subtable_title(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(text.rstrip(':'))
            r.font.size=Pt(11); r.font.name="Arial"
            r.font.bold=True; r.font.color.rgb=RGBColor(0x1B,0x2A,0x4A)

        def make_table():
            t = doc.add_table(rows=0, cols=2)
            t.style = 'Table Grid'
            try:
                tblPr = t._tbl.tblPr
                tblW  = OxmlElement('w:tblW')
                tblW.set(qn('w:w'),'9026'); tblW.set(qn('w:type'),'dxa')
                tblPr.append(tblW)
                cols_el = OxmlElement('w:tblGrid')
                for w_ in [3000, 6026]:
                    gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w_))
                    cols_el.append(gc)
                t._tbl.insert(0, cols_el)
            except: pass
            return t

        def add_table_row(table, field, value, is_header=False):
            row = table.add_row()
            trPr = row._tr.get_or_add_trPr()
            cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trPr.append(cs)
            fc = row.cells[0]; fc.text=""
            fp = fc.paragraphs[0]
            fr = fp.add_run(field)
            fr.font.size=Pt(9.5); fr.font.name="Arial"; fr.font.bold=True
            tc1 = fc._tc; tcPr1 = tc1.get_or_add_tcPr()
            shd1 = OxmlElement('w:shd')
            shd1.set(qn('w:val'),'clear'); shd1.set(qn('w:color'),'auto')
            shd1.set(qn('w:fill'),'1A5CB8' if is_header else 'E8F0FE')
            tcPr1.append(shd1)
            if is_header: fr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            for m in ['top','bottom','left','right']:
                mg = OxmlElement(f'w:{m}'); mg.set(qn('w:w'),'50'); mg.set(qn('w:type'),'dxa')
                tcPr1.append(mg)
            vc = row.cells[1]; vc.text=""
            tc2 = vc._tc; tcPr2 = tc2.get_or_add_tcPr()
            if is_header:
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'),'clear'); shd2.set(qn('w:color'),'auto')
                shd2.set(qn('w:fill'),'2E6FD4'); tcPr2.append(shd2)
            for m in ['top','bottom','left','right']:
                mg = OxmlElement(f'w:{m}'); mg.set(qn('w:w'),'50'); mg.set(qn('w:type'),'dxa')
                tcPr2.append(mg)
            value_lines = value.split('\n') if '\n' in value else [value]
            for idx_vl, vline in enumerate(value_lines):
                vp = vc.paragraphs[0] if idx_vl==0 else vc.add_paragraph()
                vr = vp.add_run(vline.strip())
                vr.font.size=Pt(9.5); vr.font.name="Arial"
                if is_header: vr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); vr.font.bold=True
            # Arabic/RTL for value cell
            if any('\u0600' <= c <= '\u06ff' for c in value):
                for vp in vc.paragraphs:
                    vp_pPr = vp._p.get_or_add_pPr()
                    bidi = OxmlElement("w:bidi"); vp_pPr.append(bidi)
                    jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); vp_pPr.append(jc)

        def add_rtl_para(text, bold=False, size=11, color=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(3)
            pPr3 = p._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi"); pPr3.append(bidi)
            jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr3.append(jc)
            r = p.add_run(text)
            r.font.size=Pt(size); r.font.name="Arial"; r.bold=bold
            if color: r.font.color.rgb=color

        # Parse report
        in_table = False
        current_table = None
        lines = rt.split('\n')
        for line in lines:
            ls = line.strip()
            if not ls:
                in_table=False; current_table=None
                doc.add_paragraph().paragraph_format.space_after=Pt(1)
                continue
            # Section title
            if re.match(r'^\d+\.\s+[A-Z\s&/]+$', ls) or ls in ('CLINICAL SUMMARY','REPORT HEADER','ORIGINAL ARABIC RESPONSES','الردود العربية الأصلية / ORIGINAL ARABIC RESPONSES'):
                in_table=False; current_table=None
                add_section_title(ls)
                continue
            # Arabic section heading
            if ls.startswith('الردود') or ('ARABIC' in ls.upper() and 'RESPONSES' in ls.upper()):
                in_table=False; current_table=None
                add_section_title(ls)
                continue
            # Separator
            if ls.startswith('━') or ls.startswith('═') or ls.startswith('===') or set(ls)=={'='}:
                in_table=False; current_table=None
                p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2)
                pPr4=p._p.get_or_add_pPr(); pBdr4=OxmlElement('w:pBdr')
                b4=OxmlElement('w:bottom'); b4.set(qn('w:val'),'single')
                b4.set(qn('w:sz'),'4'); b4.set(qn('w:space'),'1'); b4.set(qn('w:color'),'CCCCCC')
                pBdr4.append(b4); pPr4.append(pBdr4)
                continue
            # Sub-table title
            if ls.endswith(':') and '|' not in ls and len(ls)<60 and ls[0].isupper():
                in_table=False; current_table=None
                add_subtable_title(ls)
                current_table=make_table()
                add_table_row(current_table,"Field","Value",is_header=True)
                in_table=True
                continue
            # Arabic heading
            if ls.endswith(':') and any('\u0600'<=c<='\u06ff' for c in ls):
                in_table=False; current_table=None
                add_rtl_para(ls.rstrip(':'), bold=True, size=11, color=RGBColor(0x1B,0x2A,0x4A))
                continue
            # Pipe row
            if '|' in ls:
                parts=[p_.strip() for p_ in ls.split('|') if p_.strip()]
                if all(set(p_)<=set('-: ') for p_ in parts): continue
                skip_kw=[("field","value"),("milestone","finding"),("item","detail")]
                if len(parts)>=2 and (parts[0].strip('* ').lower(),parts[1].strip('* ').lower()) in skip_kw: continue
                if not in_table or current_table is None:
                    in_table=True; current_table=make_table()
                    add_table_row(current_table,"Field","Value",is_header=True)
                if len(parts)>=2:
                    add_table_row(current_table, parts[0].strip('* '), ' | '.join(parts[1:]))
                elif len(parts)==1:
                    add_table_row(current_table, parts[0].strip('* '), '')
                continue
            # Normal line
            in_table=False; current_table=None
            is_arabic=any('\u0600'<=c<='\u06ff' for c in ls)
            if is_arabic:
                add_rtl_para(ls)
            else:
                p=doc.add_paragraph()
                p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2)
                r=p.add_run(ls); r.font.size=Pt(11); r.font.name="Arial"

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    # ══ BUILD PDF ══════════════════════════════════════════
    def build_pdf(rt, pn):
        import datetime as dt
        DARK   = colors.HexColor("#1C1917")
        WARM   = colors.HexColor("#8B7355")
        BLUE   = colors.HexColor("#1A5CB8")
        BLUELT = colors.HexColor("#E8F0FE")
        BORDER = colors.HexColor("#DDD5C8")
        WHITE  = colors.white

        buf = io.BytesIO()
        doc2 = SimpleDocTemplate(buf, pagesize=A4,
                                 leftMargin=2.2*cm, rightMargin=2.2*cm,
                                 topMargin=2*cm, bottomMargin=2*cm)

        title_s   = ParagraphStyle("T",  fontName="Times-Roman",      fontSize=20, textColor=BLUE,  alignment=TA_CENTER, spaceAfter=3)
        sub_s     = ParagraphStyle("S",  fontName="Times-Italic",      fontSize=10, textColor=WARM,  alignment=TA_CENTER, spaceAfter=2)
        meta_s    = ParagraphStyle("M",  fontName="Helvetica",         fontSize=8,  textColor=WARM,  alignment=TA_CENTER, spaceAfter=12)
        section_s = ParagraphStyle("Se", fontName="Helvetica-Bold",    fontSize=10, textColor=BLUE,  spaceBefore=12, spaceAfter=4)
        body_s    = ParagraphStyle("B",  fontName="Helvetica",         fontSize=9.5,textColor=DARK,  leading=15, spaceAfter=5)
        small_s   = ParagraphStyle("Sm", fontName="Helvetica",         fontSize=8.5,textColor=WARM,  leading=13)
        arabic_s  = ParagraphStyle("Ar", fontName="Helvetica",         fontSize=10, textColor=DARK,  leading=16, spaceAfter=4, alignment=TA_RIGHT)
        footer_s  = ParagraphStyle("Ft", fontName="Helvetica-Oblique", fontSize=7.5,textColor=WARM,  leading=11, alignment=TA_CENTER)

        story=[]
        date_str = dt.datetime.now().strftime("%B %d, %Y  |  %H:%M")

        if os.path.exists(LOGO_FILE):
            try:
                logo = RLImage(LOGO_FILE, width=3.5*cm, height=1.8*cm)
                logo.hAlign="CENTER"; story.append(logo); story.append(Spacer(1,0.2*cm))
            except: pass

        story.append(Paragraph("Clinical History Report", title_s))
        story.append(Paragraph("استمارة التاريخ المرضي السريري", sub_s))
        story.append(Paragraph(f"CONFIDENTIAL  ·  {date_str}", meta_s))
        story.append(HRFlowable(width="100%", thickness=1.5, color=BLUE))
        story.append(Spacer(1, 0.3*cm))

        # Parse report lines for PDF
        in_tbl=False; tbl_data=[]

        def flush_table():
            nonlocal tbl_data, in_tbl
            if not tbl_data: return
            col_w=[4*cm, 13*cm]
            t=Table(tbl_data, colWidths=col_w)
            ts=[
                ("BACKGROUND",(0,0),(-1,0),BLUE),
                ("TEXTCOLOR",(0,0),(-1,0),WHITE),
                ("FONTNAME",(0,0),(-1,-1),"Helvetica"),
                ("FONTSIZE",(0,0),(-1,-1),8.5),
                ("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),
                ("BACKGROUND",(0,1),(-1,-1),BLUELT),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,BLUELT]),
                ("BOX",(0,0),(-1,-1),0.5,BORDER),
                ("INNERGRID",(0,0),(-1,-1),0.3,BORDER),
                ("TOPPADDING",(0,0),(-1,-1),5),
                ("BOTTOMPADDING",(0,0),(-1,-1),5),
                ("LEFTPADDING",(0,0),(-1,-1),8),
                ("VALIGN",(0,0),(-1,-1),"TOP"),
            ]
            t.setStyle(TableStyle(ts))
            story.append(t); story.append(Spacer(1,0.2*cm))
            tbl_data=[]; in_tbl=False

        for line in rt.split('\n'):
            ls=line.strip()
            if not ls:
                if in_tbl: flush_table()
                story.append(Spacer(1,0.15*cm))
                continue
            if re.match(r'^\d+\.\s+[A-Z\s&/]+$',ls) or ls in ('CLINICAL SUMMARY','REPORT HEADER') or 'ARABIC' in ls.upper():
                if in_tbl: flush_table()
                story.append(Paragraph(ls, section_s))
                story.append(HRFlowable(width="100%",thickness=0.5,color=BORDER))
                continue
            if ls.startswith('━') or ls.startswith('═') or set(ls)=={'='}:
                if in_tbl: flush_table()
                story.append(HRFlowable(width="100%",thickness=0.5,color=BORDER))
                continue
            if '|' in ls:
                parts=[p_.strip() for p_ in ls.split('|') if p_.strip()]
                if all(set(p_)<=set('-: ') for p_ in parts): continue
                skip=[("field","value"),("milestone","finding")]
                if len(parts)>=2 and (parts[0].lower().strip('* '),parts[1].lower().strip('* ')) in skip: continue
                if not in_tbl:
                    tbl_data=[[Paragraph("<b>Field</b>",small_s),Paragraph("<b>Value</b>",small_s)]]
                    in_tbl=True
                if len(parts)>=2:
                    tbl_data.append([
                        Paragraph(parts[0].strip('* '),small_s),
                        Paragraph(' | '.join(parts[1:]),body_s)
                    ])
                continue
            if in_tbl: flush_table()
            is_ar=any('\u0600'<=c<='\u06ff' for c in ls)
            if is_ar:
                story.append(Paragraph(ls, arabic_s))
            else:
                story.append(Paragraph(ls, body_s))

        if in_tbl: flush_table()

        story.append(Spacer(1,0.5*cm))
        story.append(HRFlowable(width="100%",thickness=0.5,color=BORDER))
        story.append(Spacer(1,0.2*cm))
        story.append(Paragraph(
            "This report is strictly confidential and intended solely for the treating clinician.",
            footer_s
        ))
        doc2.build(story)
        buf.seek(0)
        return buf

    # ══ ACTION BUTTONS ════════════════════════════════════
    st.markdown("<br>", unsafe_allow_html=True)
    c1,c2,c3,c4 = st.columns(4)

    with c1:
        docx_buf = build_docx(rt, pn, rs, rb_)
        st.download_button(
            "📄 Download Word",
            data=docx_buf,
            file_name=f"{fn_base}_HistoryReport.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    with c2:
        pdf_buf = build_pdf(rt, pn)
        st.download_button(
            "📑 Download PDF",
            data=pdf_buf,
            file_name=f"{fn_base}_HistoryReport.pdf",
            mime="application/pdf",
        )

    with c3:
        if st.button("📧 Send by Email"):
            try:
                docx_buf2 = build_docx(rt, pn, rs, rb_)
                pdf_buf2  = build_pdf(rt, pn)
                msg = MIMEMultipart()
                msg['From']    = GMAIL_USER
                msg['To']      = RECIPIENT_EMAIL
                msg['Subject'] = f"Clinical History Report — {pn} ({rs})"
                msg.attach(MIMEText(
                    f"Please find the clinical history report for: {pn}\n"
                    f"Form type: {rs}\nSpecialist: {rb_}\n\n"
                    f"Both Word and PDF versions are attached.",
                    'plain'
                ))
                for buf_, fname_ in [
                    (docx_buf2, f"{fn_base}_HistoryReport.docx"),
                    (pdf_buf2,  f"{fn_base}_HistoryReport.pdf"),
                ]:
                    part = MIMEBase('application','octet-stream')
                    part.set_payload(buf_.read())
                    encoders.encode_base64(part)
                    part.add_header('Content-Disposition',f'attachment; filename="{fname_}"')
                    msg.attach(part)
                with smtplib.SMTP_SSL('smtp.gmail.com',465) as server:
                    server.login(GMAIL_USER, GMAIL_PASS)
                    server.sendmail(GMAIL_USER, RECIPIENT_EMAIL, msg.as_string())
                st.success(f"✅ Sent to {RECIPIENT_EMAIL}")
            except Exception as e:
                st.error(f"Email error: {str(e)}")

    with c4:
        if st.button("↺ New Patient"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
